# -*- coding: utf-8 -*-
"""
EzSell FYP Final Report Builder - Part 1: Introduction Chapter
Fills the template starting at paragraph index 180 (page 7)
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

TEMPLATE = r'c:\Users\ahmed\ezsell\FYP_Final Report_Sring 2026_Template (Updated) (1).docx'
OUTPUT   = r'c:\Users\ahmed\ezsell\EZSell_FYP_Final_Report.docx'

# ─── helpers ────────────────────────────────────────────────────────────────
def cell_shade(cell, hex_col):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_col)
    tcPr.append(shd)

def h1(doc, txt):  return doc.add_heading(txt, level=1)
def h2(doc, txt):  return doc.add_heading(txt, level=2)
def h3(doc, txt):  return doc.add_heading(txt, level=3)
def para(doc, txt, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(txt)
    return p
def bpara(doc, txt):
    p = doc.add_paragraph()
    p.add_run(txt).bold = True
    return p
def bullet(doc, txt):
    p = doc.add_paragraph(style='List Paragraph')
    p.add_run(txt)
    p.paragraph_format.left_indent = Inches(0.25)
    return p

def add_table_caption(doc, txt):
    p = doc.add_paragraph(txt, style='Caption')
    return p

def simple_table(doc, header_row, rows, header_color='1F3864', header_text_color='FFFFFF'):
    """Add a table with a header and data rows."""
    t = doc.add_table(rows=1+len(rows), cols=len(header_row))
    t.style = 'Table Grid'
    # header
    hrow = t.rows[0]
    for i, h in enumerate(header_row):
        c = hrow.cells[i]
        c.text = h
        cell_shade(c, header_color)
        for run in c.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    # data
    alt = 'DCE6F1'
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            if ri % 2 == 0:
                cell_shade(row.cells[ci], alt)
            for run in row.cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9)
    return t

def algo_table(doc, title, inputs, outputs, steps):
    """Pseudocode-style algorithm table (3-column merged header)."""
    ncols = 3
    nrows = 3 + len(steps)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = 'Table Grid'
    # merge header across all cols
    t.rows[0].cells[0].merge(t.rows[0].cells[ncols-1])
    c = t.rows[0].cells[0]
    c.text = title
    cell_shade(c, '1F3864')
    for run in c.paragraphs[0].runs:
        run.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(10)
    # inputs row
    t.rows[1].cells[0].merge(t.rows[1].cells[ncols-1])
    t.rows[1].cells[0].text = 'Input: ' + inputs
    # outputs row
    t.rows[2].cells[0].merge(t.rows[2].cells[ncols-1])
    t.rows[2].cells[0].text = 'Output: ' + outputs
    # steps
    for i, step in enumerate(steps):
        t.rows[3+i].cells[0].merge(t.rows[3+i].cells[ncols-1])
        t.rows[3+i].cells[0].text = step
    return t

# ─── Load template and truncate at intro chapter (index 180) ────────────────
doc = Document(TEMPLATE)
body = doc.element.body
children = list(body)

para_count = 0
cut_idx = None
for i, child in enumerate(children):
    if child.tag.endswith('}p') or child.tag.endswith('}tbl') or child.tag.endswith('}sectPr'):
        if child.tag.endswith('}p'):
            if para_count == 180:
                cut_idx = i
                break
            para_count += 1

if cut_idx is not None:
    # Remove elements from cut_idx onward, but preserve the last sectPr
    to_remove = list(children)[cut_idx:]
    # Separate sectPr elements from the rest
    secPrs = [el for el in to_remove if el.tag.endswith('}sectPr')]
    for el in to_remove:
        body.remove(el)
    # Re-attach the last sectPr so section info is preserved for table width
    if secPrs:
        body.append(secPrs[-1])
    print(f"Truncated at element index {cut_idx}, re-attached sectPr: {bool(secPrs)}")
else:
    print("WARNING: Cut point not found, appending to end")

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '1. Introduction')
para(doc,
    'This chapter provides an overview of the EzSell project, including its vision, '
    'a comparative analysis of existing marketplace systems, key deliverables, system '
    'constraints, technology stack, and relevance to the academic curriculum. EzSell '
    'is a full-stack, AI-powered second-hand marketplace platform tailored for the '
    'Pakistani market, enabling intelligent buying and selling of used mobile phones, '
    'laptops, and furniture through a secure, visually rich web application.')

# 1.1 Vision Statement
h2(doc, '1.1 Vision Statement')
para(doc,
    'For individual buyers and sellers in Pakistan who face challenges in pricing '
    'second-hand electronics and furniture accurately and safely, EzSell is an '
    'AI-powered web marketplace that provides real-time price intelligence, '
    'automated fraud prevention, AR visualization, and a personalized buying '
    'experience. Unlike existing platforms such as OLX and Facebook Marketplace, '
    'which lack AI-assisted pricing, structured fraud detection, and immersive '
    'product visualization, EzSell leverages a multi-model ML ensemble, a Groq '
    'LLM layer, CLIP-based image validation, SentenceTransformer semantic '
    'recommendations, and Tripo AI 3D generation to deliver a trustworthy, '
    'data-driven, and visually superior marketplace for the pre-owned goods economy '
    'of Pakistan.')

# 1.2 Related System Analysis
h2(doc, '1.2 Related System Analysis / Literature Review')
para(doc,
    'The following table presents a comparative analysis of three leading marketplace '
    'platforms against the proposed EzSell solution, identifying their weaknesses and '
    'the specific innovations EzSell introduces to address each gap.')

add_table_caption(doc, 'Table 1   Related System Analysis with Proposed Project Solution')
simple_table(doc,
    ['Application Name', 'Weakness', 'Proposed Project Solution'],
    [
        ['OLX Pakistan',
         'No AI-assisted pricing; prices are entirely seller-determined leading to '
         'widespread overpricing and underpricing. High incidence of scam listings '
         'with no automated image or content verification.',
         'EzSell integrates a dual-track ML + LLM pricing pipeline (XGBoost, '
         'LightGBM, Groq llama-3.3-70b-versatile) with live OLX market scraping '
         'and IQR outlier filtering. CLIP-based image validation and perceptual '
         'hashing eliminate fraudulent listings before they go live.'],
        ['Facebook Marketplace',
         'Lacks structured pricing guidance; relies solely on social trust signals. '
         'No category-specific fields for electronics. Buyers cannot visualize '
         'furniture in their actual space before purchase.',
         'EzSell provides category-specific attribute fields (RAM, storage, '
         'processor for mobiles/laptops; material, dimensions for furniture), '
         'AI price suggestions, and a full AR Try-On viewer powered by Tripo AI '
         'and Google model-viewer for in-room furniture visualization.'],
        ['eBay',
         'Complex auction-based process not suited for local, face-to-face '
         'Pakistani commerce. High commission fees deter casual sellers. '
         'No localized market data for Pakistani price ranges.',
         'EzSell is a free-to-list, direct-sale platform trained on Pakistani '
         'market data scraped from OLX Pakistan and retail tech sites. '
         'All price predictions reflect PKR-denominated local market realities '
         'using IQR-filtered datasets of thousands of live listings.'],
        ['Daraz.pk',
         'Primarily a new-goods B2C platform; does not support C2C used-goods '
         'listings. Sellers require business registration. No secondhand '
         'electronics or furniture ecosystem.',
         'EzSell is purpose-built for the secondhand C2C market, supporting '
         'individual sellers with full listing lifecycle management (create, '
         'edit, hide, mark-as-sold) and AI-assisted pricing aligned to '
         'real-world used-goods valuations.'],
    ]
)

# 1.3 Project Deliverables
h2(doc, '1.3 Project Deliverables')
para(doc, 'The following deliverables constitute the complete EzSell 1.0.0 production release:')
deliverables = [
    'D-1: Fully functional FastAPI backend with RESTful APIs for all marketplace operations.',
    'D-2: React 18 + Vite + TypeScript frontend SPA with Tailwind CSS and Shadcn UI component library.',
    'D-3: AI-powered price prediction engine — trained ML ensemble (.pkl models) per category with Groq LLM validation layer.',
    'D-4: Automated fraud prevention pipeline — CLIP image-category validation, dHash duplicate detection, scam keyword filter.',
    'D-5: Semantic recommendation engine — SentenceTransformer (all-MiniLM-L6-v2) embedding-based personalized listing feed.',
    'D-6: AR/3D furniture try-on module — Tripo AI GLB generation, WebAR viewer, iOS USDZ QuickLook export.',
    'D-7: Admin analytical dashboard with listing moderation, user management, support ticket system, and engagement metrics.',
    'D-8: Real-time in-app messaging system with unread-count badges and conversation inbox.',
    'D-9: AI chatbot assistant (EzSell Assistant) using Groq llama-3.1-8b-instant with OLX CSV price context injection.',
    'D-10: Scraped and preprocessed market datasets for mobile phones, laptops, and furniture (OLX Pakistan).',
    'D-11: Cloudinary-integrated cloud image storage with local fallback.',
    'D-12: Google OAuth 2.0 authentication alongside JWT-based session management and email OTP verification.',
]
for d in deliverables:
    bullet(doc, d)

# 1.4 System Limitations/Constraints
h2(doc, '1.4 System Limitations / Constraints')
limits = [
    'LI-1: EzSell is a web-only platform at launch; no dedicated iOS or Android native application is provided, though the frontend is fully mobile-responsive.',
    'LI-2: Only three product categories are supported in v1.0.0 — Mobiles, Laptops, and Furniture. Vehicles, fashion, and other categories are deferred to future releases.',
    'LI-3: Price prediction accuracy is dependent on the quality and volume of scraped OLX data; very new or ultra-niche models with fewer than 5 market data points may yield lower-confidence predictions.',
    'LI-4: Web scraping of OLX Pakistan via DuckDuckGo DDGS is subject to external rate limits and website structural changes; the platform provides IQR-filtered cached CSV data as a fallback.',
    'LI-5: The platform does not facilitate in-app payments or escrow; all financial transactions are completed offline between buyer and seller.',
    'LI-6: AR 3D generation via Tripo AI requires a paid API key with credit balance; generation takes 30–90 seconds and depends on third-party API availability.',
    'LI-7: Email verification is required before a user can post listings; phone number is collected but biometric verification is not implemented in v1.0.0.',
    'LI-8: AR advanced features (AI-generated 3D) require devices with dedicated GPU support (flagship smartphones 2021+, or desktop with hardware acceleration).',
    'LI-9: The chatbot (EzSell Assistant) is scoped exclusively to EzSell-related queries; general-purpose AI assistance is deliberately restricted by guardrails.',
]
for l in limits:
    para(doc, l)

# 1.5 Tools and Technologies
h2(doc, '1.5 Tools and Technologies')
para(doc,
    'Table 2 below lists all hardware and software tools, frameworks, libraries, '
    'APIs, and services used in the implementation of EzSell 1.0.0.')

add_table_caption(doc, 'Table 2   Tools and Technologies for EzSell')
simple_table(doc,
    ['Category', 'Tool / Technology', 'Version', 'Rationale'],
    [
        ['Programming Language', 'Python', '3.11+', 'Backend logic, ML pipeline, AI services, web scraping'],
        ['Programming Language', 'TypeScript', '5.8', 'Type-safe frontend development with React'],
        ['Programming Language', 'JavaScript (ES2023)', 'ES2023', 'Frontend runtime and browser interactions'],
        ['Frontend Framework', 'React', '18.3', 'Component-based SPA architecture with hooks-based state management'],
        ['Frontend Build Tool', 'Vite', '5.4', 'Ultra-fast HMR dev server and production bundler with SWC plugin'],
        ['UI Component Library', 'Shadcn UI + Radix UI', 'Latest', 'Accessible, unstyled headless UI primitives composed with Tailwind'],
        ['Styling', 'Tailwind CSS', '3.4', 'Utility-first CSS framework for rapid, consistent UI development'],
        ['3D / AR', 'Google model-viewer', '4.1', 'Web Component for WebAR GLB rendering with iOS USDZ QuickLook'],
        ['3D / AR', '@react-three/fiber + drei', '8.18 / 9.122', 'React renderer for Three.js; used for procedural 3D scenes'],
        ['3D / AR', 'Three.js', '0.182', 'Core 3D engine for WebGL rendering and GLB/GLTF loading'],
        ['Backend Framework', 'FastAPI', '0.115+', 'High-performance async Python REST API with automatic OpenAPI docs'],
        ['ASGI Server', 'Uvicorn', 'Latest', 'Production ASGI server for FastAPI with Gunicorn process management'],
        ['ORM', 'SQLAlchemy', '2.x', 'Database-agnostic ORM supporting SQLite (dev) and PostgreSQL (prod)'],
        ['Data Validation', 'Pydantic v2', '2.x', 'Schema validation, serialization, and settings management'],
        ['Database', 'SQLite', '3.x', 'Development database; file-based with zero configuration'],
        ['Database', 'PostgreSQL', '15', 'Production-grade relational database for deployment'],
        ['Authentication', 'python-jose + passlib', 'Latest', 'JWT token generation/verification; bcrypt password hashing'],
        ['Authentication', 'Google OAuth 2.0', 'v2', 'Social login via Google; OAuth authorization code flow'],
        ['ML — Pricing', 'XGBoost', '2.x', 'Primary gradient boosted ensemble model for price prediction (35% weight)'],
        ['ML — Pricing', 'LightGBM', 'Latest', 'Fast gradient boosting with leaf-wise tree growth (35% weight)'],
        ['ML — Pricing', 'Scikit-learn', '1.x', 'Random Forest (15%) + Gradient Boosting (15%); preprocessing pipelines'],
        ['ML — Pricing', 'Joblib', 'Latest', 'Serialization and loading of trained .pkl model artifacts'],
        ['ML — Embeddings', 'SentenceTransformers', '3.x', 'Semantic embedding model (all-MiniLM-L6-v2) for recommendations'],
        ['LLM — Pricing/Chat', 'Groq API', 'Latest', 'llama-3.3-70b-versatile (pricing); llama-3.1-8b-instant (chatbot)'],
        ['Vision — Fraud', 'OpenAI CLIP (ViT-B/32)', 'Latest', 'Zero-shot image-to-category classification for listing validation'],
        ['Vision — Fraud', 'Pillow (PIL)', '10.x', 'dHash perceptual image hashing for duplicate photo detection'],
        ['AR Generation', 'Tripo AI API', 'V2', 'Image-to-3D GLB model generation with Draco mesh optimization'],
        ['Web Scraping', 'DuckDuckGo DDGS', 'Latest', 'Search-based OLX Pakistan price scraping for live market data'],
        ['Data Processing', 'Pandas', '2.x', 'CSV ingestion, IQR filtering, market data preprocessing'],
        ['Data Processing', 'NumPy', '1.x', 'Numerical operations for pricing statistics and vector math'],
        ['Image Storage', 'Cloudinary', 'Latest', 'Cloud image CDN with transformation APIs; local filesystem fallback'],
        ['Email', 'SMTP (smtplib)', 'Python std', 'OTP verification emails and password reset link delivery'],
        ['HTTP Client', 'Axios', '1.13', 'Frontend REST API client with interceptors and request cancellation'],
        ['State Management', 'TanStack React Query', '5.90', 'Server-state caching, background refetching, and mutation handling'],
        ['Forms', 'React Hook Form + Zod', '7.61 / 3.25', 'Performant form state management with schema-based validation'],
        ['Charts', 'Recharts', '2.15', 'Declarative chart library for admin analytics dashboards'],
        ['Version Control', 'Git + GitHub', 'Latest', 'Source code versioning and collaborative development'],
        ['API Testing', 'Postman', 'Latest', 'REST API endpoint testing and documentation'],
        ['IDE', 'Visual Studio Code', 'Latest', 'Primary code editor with TypeScript/Python extension support'],
        ['Deployment', 'Cloudinary + Local FS', 'Latest', 'Dual-track media storage with aggressive CDN caching (1-year max-age)'],
    ]
)

# 1.6 Relevance to Course Modules
h2(doc, '1.6 Relevance to Course Modules')
para(doc,
    'EzSell integrates knowledge from multiple Computer Science core and elective '
    'modules studied throughout the BCS programme at AIR University, Islamabad:')
courses = [
    ('Software Engineering (SE)', 'Requirements analysis (SRS), system design (SDS), testing strategies (unit, functional, integration), and SDLC management practices directly applied throughout the project lifecycle.'),
    ('Database Systems (DB)', 'SQLAlchemy ORM design, relational schema normalization, foreign key relationships, query optimization, and migration management for User, Listing, Message, Notification, and analytics tables.'),
    ('Artificial Intelligence (AI)', 'ML model selection (XGBoost, LightGBM, Random Forest, Gradient Boosting), feature engineering from structured and unstructured listing data, and ensemble learning strategies.'),
    ('Machine Learning (ML)', 'Price prediction model training pipeline including hyperparameter tuning (RandomizedSearchCV), cross-validation (5-fold CV), IQR statistical filtering, and model evaluation metrics (R², MAE, RMSE, MAPE).'),
    ('Natural Language Processing (NLP)', 'Keyword extraction for recommendation interests, SentenceTransformer semantic embeddings, cosine similarity-based listing ranking, and LLM-based title validation via Groq API.'),
    ('Computer Vision (CV)', 'CLIP zero-shot image classification for fraud detection, dHash perceptual hashing for duplicate image detection, and OpenCV-based image analysis.'),
    ('Web Technologies', 'Full-stack development: FastAPI RESTful backend, React 18 SPA with Vite, Tailwind CSS styling, and browser-native WebAR with model-viewer.'),
    ('Computer Networks', 'REST API communication, HTTP/HTTPS protocols, CORS middleware configuration, SSE (Server-Sent Events) for chatbot streaming, and CDN caching strategies.'),
    ('Human-Computer Interaction (HCI)', 'UI/UX design principles applied in the listing creation flow, AR viewer coaching overlays, chatbot personality design, and responsive mobile-first layouts.'),
    ('Information Security', 'JWT authentication, bcrypt password hashing, Google OAuth 2.0, email OTP verification, and role-based access control (admin vs. regular user).'),
    ('3D Computer Graphics', 'Three.js/WebGL 3D scene construction, GLB/GLTF model loading, USDZ export for iOS AR QuickLook, and Draco mesh compression for performance optimization.'),
]
for course, desc in courses:
    p = doc.add_paragraph(style='List Paragraph')
    p.add_run(course + ': ').bold = True
    p.add_run(desc)

doc.save(OUTPUT)
print(f"\nPart 1 complete. Saved: {OUTPUT}")
