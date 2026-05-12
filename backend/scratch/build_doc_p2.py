# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'c:\Users\ahmed\ezsell\EZSell_FYP_Final_Report.docx'
doc = Document(OUTPUT)

def cell_shade(cell, hex_col):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_col); tcPr.append(shd)

def h1(doc,t): return doc.add_heading(t,level=1)
def h2(doc,t): return doc.add_heading(t,level=2)
def h3(doc,t): return doc.add_heading(t,level=3)
def para(doc,t): p=doc.add_paragraph(); p.add_run(t); return p
def bul(doc,t):
    p=doc.add_paragraph(style='List Paragraph')
    p.add_run(t); p.paragraph_format.left_indent=Inches(0.25); return p

def simple_table(doc, headers, rows, hcol='1F3864'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; cell_shade(c,hcol)
        for r in c.paragraphs[0].runs:
            r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]
        for ci,val in enumerate(row):
            tr.cells[ci].text=val
            if ri%2==0: cell_shade(tr.cells[ci],'DCE6F1')
            for r in tr.cells[ci].paragraphs[0].runs: r.font.size=Pt(9)
    return t

def caption(doc,t): return doc.add_paragraph(t,style='Caption')

# ══════════════════════════════════════════════════════
# CHAPTER 2 – PROBLEM DEFINITION
# ══════════════════════════════════════════════════════
doc.add_page_break()
h1(doc,'2. Problem Definition')
para(doc,
 'This chapter defines the precise problem that EzSell addresses, articulates '
 'the proposed software solution, lists measurable business objectives, and '
 'defines the scope of the system. It also enumerates the functional modules '
 'that collectively form the EzSell 1.0.0 product.')

h2(doc,'2.1 Problem Statement')
para(doc,
 'The Pakistani second-hand goods market — particularly for mobile phones, laptops, '
 'and furniture — is one of the fastest-growing informal economies in South Asia, '
 'yet it remains riddled with systemic inefficiencies that harm both buyers and sellers '
 'from an economic and psychological perspective. Existing platforms such as OLX Pakistan '
 'operate as unstructured classified advertisement boards with no artificial intelligence '
 'integration, no automated pricing guidance, and no verifiable trust mechanisms.')
para(doc,
 'Psychologically, the absence of pricing transparency induces high cognitive load and '
 'decision fatigue. Sellers routinely set prices based purely on intuition or emotional '
 'attachment to the item, leading to either significant under-pricing (financial loss) '
 'or over-pricing (prolonged listing periods and anxiety over failed sales). Buyers, on '
 'the other hand, experience a chronic "trust deficit" and "buyer\'s remorse" anxiety. '
 'The absence of reliable pricing benchmarks forces them to negotiate blindly, making '
 'them vulnerable to inflated prices.')
para(doc,
 'Furthermore, fraudulent listings — including duplicate advertisements, stolen product '
 'photos, and scam-baiting descriptions (e.g., "advance payment required") — exploit this '
 'lack of structural trust, creating an environment of perpetual suspicion. For furniture '
 'buyers specifically, the inability to visualize a piece of furniture in their actual '
 'living space before purchase leads to spatial anxiety and costly mismatches in size and style.')
para(doc,
 'The absence of a semantic recommendation engine means buyers must manually filter '
 'through hundreds of irrelevant listings, leading to search fatigue. In summary, '
 'the Pakistani second-hand marketplace ecosystem lacks the AI-powered price prediction, '
 'automated multi-layer fraud detection, and immersive AR visualization required to '
 'transform a high-anxiety, low-trust environment into a secure, structured, and '
 'psychologically reassuring digital marketplace.')

h2(doc,'2.2 Problem Solution')
para(doc,
 'EzSell is designed as a comprehensive answer to every identified problem in the '
 'Pakistani second-hand marketplace. The platform is built on a FastAPI backend '
 'and React 18 frontend, integrating a suite of AI, ML, and computer vision '
 'technologies to deliver the following solutions:')
solutions = [
 ('AI-Powered Price Intelligence',
  'A dual-track pricing pipeline combines a trained ML ensemble (XGBoost 35%, '
  'LightGBM 35%, Random Forest 15%, Gradient Boosting 15%) with a Groq LLM layer '
  '(llama-3.3-70b-versatile) that performs live OLX Pakistan market scraping. '
  'IQR statistical filtering removes fake outlier listings, producing a confident, '
  'market-grounded price estimate with an 85%+ accuracy rate within a ±20% price margin.'),
 ('Multi-Layer Fraud Prevention',
  'Every submitted listing passes through a 5-stage fraud pipeline: (1) email '
  'verification gate, (2) content hash deduplication for copy-paste ads, '
  '(3) CLIP-based image-to-category AI validation, (4) dHash perceptual hashing '
  'for stolen-photo detection, and (5) scam keyword scanning. Flagged listings '
  'enter a pending-review queue before going live.'),
 ('Augmented Reality Furniture Try-On',
  'EzSell integrates the Tripo AI V2 API to generate photorealistic GLB 3D models '
  'from listing images. Buyers can view furniture in their room via a WebAR browser '
  'session (Google model-viewer) or natively on iOS via USDZ QuickLook export, '
  'with real-world-accurate dimensions locked in cm.'),
 ('Semantic Recommendation Engine',
  'SentenceTransformer (all-MiniLM-L6-v2) encodes listing content into 384-dimensional '
  'semantic vectors. User interest profiles are built from behavioral signals '
  '(views, searches, favorites, messages) with time-decay weighting. Cosine similarity '
  'between user interest embeddings and listing vectors powers a personalized feed.'),
 ('Structured Marketplace with Category Intelligence',
  'EzSell enforces structured category-specific attribute fields: RAM/storage/processor '
  'for mobiles and laptops; furniture type/material/subtype for furniture. An NLP '
  'keyword extractor and regex-based furniture subtype inferencer ensure rich, '
  'machine-readable listing metadata.'),
 ('Real-Time Messaging and Trust Signals',
  'Buyers can message sellers directly through a conversation inbox tied to specific '
  'listings. Seller phone numbers are shown on listing detail pages. Profile trust '
  'signals (verified email badge, rating, listing history) create accountability.'),
 ('AI Chatbot Assistant',
  'An in-app streaming chatbot powered by Groq llama-3.1-8b-instant provides '
  'contextual assistance for navigation, listing help, AR guidance, fraud '
  'explanation, price intelligence, and furniture style advice — injecting '
  'live OLX CSV price context for pricing queries.'),
]
for title, desc in solutions:
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = Inches(0.25)
    p.add_run(title + ': ').bold = True
    p.add_run(desc)

h2(doc,'2.3 Objectives of the Proposed System')
objectives = [
 'BO-1: Deliver AI-assisted price prediction achieving 85%+ accuracy within ±20% of actual market prices for mobile phones, laptops, and furniture, using a weighted ML ensemble validated by a live Groq LLM market scraping layer.',
 'BO-2: Reduce fraudulent listing submissions by a projected 70%+ through a five-stage automated fraud prevention pipeline encompassing email verification, content hashing, CLIP image validation, perceptual hashing, and scam keyword filtering.',
 'BO-3: Enable furniture buyers to visualize products in their actual environment via browser-native WebAR (GLB) and iOS AR QuickLook (USDZ), powered by Tripo AI 3D generation from listing photographs.',
 'BO-4: Provide personalized product recommendations using SentenceTransformer semantic embeddings and behavioral user activity profiling, improving listing discovery relevance for authenticated users.',
 'BO-5: Implement secure buyer-seller communication through an in-app messaging inbox with unread-count badges, eliminating the need for unsafe external contact channels.',
 'BO-6: Provide sellers with an intelligent listing dashboard showing active, hidden, sold, and pending ad counts alongside AI-suggested pricing, reducing time-to-correct-price for new listings.',
 'BO-7: Enable administrators to manage the full platform lifecycle — listing approval/rejection with documented reasons, user account management, support ticket resolution, and engagement analytics — through a dedicated admin dashboard.',
 'BO-8: Streamline the listing creation process with category-specific attribute fields, real-time CLIP image validation, AI title validation, and an AI price suggestion widget, reducing listing errors and improving data quality.',
 'BO-9: Deliver a fully responsive, mobile-first web application that functions seamlessly across desktop, tablet, and smartphone browsers without a dedicated native application.',
 'BO-10: Provide an embedded AI assistant (EzSell Chatbot) capable of answering marketplace queries, guiding AR usage, explaining pricing, and advising on furniture styles — strictly scoped to EzSell-related topics.',
]
for obj in objectives:
    bul(doc, obj)

h2(doc,'2.4 Scope')
para(doc,
 'EzSell is a web-based, AI-augmented C2C (consumer-to-consumer) marketplace platform '
 'serving the Pakistani second-hand goods economy. The system supports three product '
 'categories at launch: Mobile Phones, Laptops, and Furniture. The platform encompasses '
 'the complete buying and selling lifecycle — from account registration and listing '
 'creation to AI-validated publishing, buyer discovery, AR visualization, in-app '
 'messaging, and post-transaction review.')
para(doc,
 'On the seller side, the system provides an AI price prediction engine that delivers '
 'market-grounded PKR price estimates derived from live OLX Pakistan scraping and a '
 'trained ML ensemble. Sellers benefit from real-time fraud screening at submission '
 'time, a personal dashboard for listing lifecycle management, and automated '
 'notifications for admin review decisions.')
para(doc,
 'On the buyer side, the system provides a personalized homepage feed ranked by '
 'semantic similarity to the user\'s inferred interests, advanced multi-parameter '
 'filtering (category, price range, brand, condition, location), and an AR '
 'try-on experience for furniture listings. Buyers can save listings to a '
 'favorites collection, message sellers directly, and view seller trust profiles.')
para(doc,
 'Administratively, the platform includes a full admin dashboard with listing '
 'moderation controls (approve/reject with reasons), user account management, '
 'category analytics charts, engagement metric tracking, and a support ticket '
 'system with automated notification dispatch.')
para(doc,
 'The scope explicitly excludes: in-app payment processing (all financial '
 'transactions occur offline), vehicle and fashion categories (deferred to v2.0), '
 'native iOS/Android mobile applications (the web app is mobile-responsive), '
 'and biometric user verification (email OTP and phone collection are implemented; '
 'biometrics are deferred). The platform is scoped to the Pakistani market with '
 'pricing, currency (PKR), and market data calibrated to local conditions.')

h2(doc,'2.5 Modules')
para(doc,'EzSell 1.0.0 comprises ten functional modules:')

modules = [
 ('Module 1: User Management & Authentication',
  ['FE-1: User registration with username, email, mandatory phone number, and bcrypt-hashed password.',
   'FE-2: Six-digit email OTP verification required before a user can post listings.',
   'FE-3: Google OAuth 2.0 social login with profile completion flow for new Google users.',
   'FE-4: JWT Bearer token authentication for all secured API endpoints (python-jose, HS256).',
   'FE-5: Password reset via email link with time-limited secure token.',
   'FE-6: Profile management — avatar upload (Cloudinary), full name, bio, location, phone.',
   'FE-7: Admin role with elevated privileges: listing moderation, user management, ticket resolution.']),
 ('Module 2: Product Listings & Marketplace',
  ['FE-1: Create listings with title, description, price, category, condition, location, and up to 5 images.',
   'FE-2: Category-specific attribute fields: Brand/Model for mobiles & laptops; Furniture Type/Material/Subtype/Color for furniture.',
   'FE-3: Regex-based furniture subtype auto-inference from listing title and description text.',
   'FE-4: Edit listing — update any field; owner-only access enforced.',
   'FE-5: Delete listing — permanent removal by owner or admin.',
   'FE-6: Toggle visibility (hide/unhide) without deleting the listing.',
   'FE-7: Mark as Sold — sets is_sold=True and auto-hides the listing.',
   'FE-8: Approval gate — listings pass fraud checks and enter pending/approved/rejected states.',
   'FE-9: Automatic 30-day expiry — listings older than 30 days are auto-deactivated.',
   'FE-10: View count increment on each product detail page visit.']),
 ('Module 3: AI-Powered Price Prediction',
  ['FE-1: Weighted ensemble ML model per category (.pkl): XGBoost (35%) + LightGBM (35%) + Random Forest (15%) + Gradient Boosting (15%).',
   'FE-2: Groq LLM layer (llama-3.3-70b-versatile) performs live OLX Pakistan scraping via DuckDuckGo DDGS.',
   'FE-3: IQR outlier filtering on scraped price sets to eliminate fake/spam listings.',
   'FE-4: NLP feature extraction: RAM, storage, processor generation, brand, condition parsed from listing text.',
   'FE-5: Dual-track output: ML ensemble price (85% weight) blended with LLM market estimate (15% weight).',
   'FE-6: Confidence score (0–100%) returned alongside the price prediction.',
   'FE-7: Predicted vs. asked price comparison shown to sellers with low-confidence warnings.',
   'FE-8: GSMArena spec scraping as a fallback for mobile phone specification extraction.']),
 ('Module 4: Smart Recommendations & Search',
  ['FE-1: User activity tracking — views, searches, favorites, clicks, messages logged to UserActivity table.',
   'FE-2: Interest aggregation with time-decay weighting: recent activities weighted higher (exp(-lambda*days)).',
   'FE-3: SentenceTransformer (all-MiniLM-L6-v2) generates 384-dim embeddings for listings and user interest profiles.',
   'FE-4: Cosine similarity ranking of listings against user embedding for personalized homepage feed.',
   'FE-5: Trending/popular listings served to anonymous (unauthenticated) visitors.',
   'FE-6: Multi-parameter search: tokenized keyword matching across title, description, brand, category, furniture type, location with stemming and stopword filtering.',
   'FE-7: Filter controls: category, price range, condition, location — server-side SQLAlchemy query composition.']),
 ('Module 5: Augmented Reality (AR) Try-On',
  ['FE-1: Procedural 3D furniture model generation from listing attributes (type, subtype, material, dimensions) using FurnitureGLBGenerator.',
   'FE-2: Tripo AI V2 API integration — image-to-3D GLB generation from listing photographs with Draco mesh compression.',
   'FE-3: WebAR viewer using Google model-viewer web component; works on modern browsers without app installation.',
   'FE-4: iOS AR QuickLook USDZ export for native iPhone/iPad AR experience.',
   'FE-5: Real-world accurate dimensions locked in centimetres per furniture subtype.',
   'FE-6: Floor shadow decal for realistic ground plane anchoring.',
   'FE-7: AR coaching overlay guides user to scan flat surface before placing furniture.',
   'FE-8: Admin manual GLB upload — admin can attach a custom GLB file to any listing.']),
 ('Module 6: Admin Analytical Dashboard',
  ['FE-1: Overview metrics: total users, total listings, active listings, pending reviews, support tickets.',
   'FE-2: Listing moderation: approve or reject pending listings with documented rejection reason stored in DB.',
   'FE-3: User management: view full user profiles, listing counts, and account status.',
   'FE-4: Category analytics: chart showing listing breakdown by category and approval status.',
   'FE-5: Engagement metrics: views, favorites, message counts per listing.',
   'FE-6: Support ticket dashboard: view all tickets with user info, type, subject, and description.',
   'FE-7: Ticket status controls: Open → Working → Done; each status change triggers user notification.']),
 ('Module 7: Fraud Prevention',
  ['FE-1: Email verification gate — unverified users cannot submit listings (HTTP 403).',
   'FE-2: Content hash deduplication — MD5 hash of title+description+price+owner_id detects copy-paste duplicates.',
   'FE-3: CLIP (ViT-B/32) zero-shot image-to-category validation — images not matching the declared category cause a hard rejection with image deletion.',
   'FE-4: dHash perceptual hashing — reused or stolen listing photos detected across all listings in the database.',
   'FE-5: Price anomaly check — price < 40% or > 300% of predicted market price flags the listing as suspicious.',
   'FE-6: Scam keyword scanning — regex patterns matching "advance payment", "WhatsApp only", "gift card" etc. flag listings for review.',
   'FE-7: Nonsense/gibberish detection — keyboard mashing, extreme character repetition, and low word diversity trigger quality flags.',
   'FE-8: Fraud flags stored as JSON array per listing; surfaced to admin in the moderation queue.']),
 ('Module 8: Messaging & Inbox',
  ['FE-1: Buyer-to-seller messaging tied to specific listings from the product detail page.',
   'FE-2: Conversation inbox grouped by listing and contact — full message thread view.',
   'FE-3: Unread message count badge in the navigation bar, polled periodically.',
   'FE-4: Opening a conversation marks all messages in it as read.',
   'FE-5: Seller phone number shown on listing detail page alongside the message button.',
   'FE-6: Auto-message generated when a buyer views seller phone number (call intent tracking).']),
 ('Module 9: Favorites & User Dashboard',
  ['FE-1: Per-user saved listings with instant heart-toggle add/remove from card or detail page.',
   'FE-2: Dedicated /favorites page showing all saved listings with remove option.',
   'FE-3: Seller personal dashboard at /dashboard showing all ads grouped by status (Active, Hidden, Sold, Pending).',
   'FE-4: Dashboard status cards: live counts of Active, Hidden, Sold, and Pending listings.',
   'FE-5: Listing analytics per ad: view count, favorites count, message count.',
para(doc,'EzSell 1.0.0 comprises eleven functional modules:')

h3(doc, 'Module 1: User Authentication & Account Management')
bul(doc, 'FE-1: Register and login securely using Email/Password with OTP verification.')
bul(doc, 'FE-2: Authenticate instantly using Google OAuth 2.0 integration.')
bul(doc, 'FE-3: Manage active stateless sessions securely via JSON Web Tokens (JWT).')

h3(doc, 'Module 2: Product Listings & Ad Management')
bul(doc, 'FE-1: Create new used-goods listings using category-specific dynamic forms for Mobiles, Laptops, and Furniture.')
bul(doc, 'FE-2: Auto-extract and populate metadata specifications using regex-based text analysis on listing titles.')
bul(doc, 'FE-3: Transition listing statuses through a complete lifecycle including Draft, Active, Hidden, and Sold.')

h3(doc, 'Module 3: Price Prediction')
bul(doc, 'FE-1: Generate a baseline market price utilizing a pre-trained ML ensemble of XGBoost, LightGBM, Random Forest, and Gradient Boosting.')
bul(doc, 'FE-2: Cross-validate baseline predictions with live web scraping of OLX Pakistan via the Groq LLM.')
bul(doc, 'FE-3: Filter scraped pricing data using Interquartile Range (IQR) mathematics to discard extreme outliers and ensure accuracy.')

h3(doc, 'Module 4: Recommendations')
bul(doc, 'FE-1: Track user behavioral metrics including views, clicks, and favorites across different listing categories.')
bul(doc, 'FE-2: Generate a personalized discovery feed by utilizing SentenceTransformer embeddings ranked via cosine similarity algorithms.')

h3(doc, 'Module 5: Analytical Dashboard')
bul(doc, 'FE-1: Provide administrators with Recharts-based visual analytics tracking platform engagement and total active users.')
bul(doc, 'FE-2: Manage a centralized moderation queue allowing admins to manually review, approve, or reject flagged user listings.')

h3(doc, 'Module 6: AR Customization Try-On')
bul(doc, 'FE-1: Convert standard 2D user-uploaded furniture images into photorealistic 3D GLB models using the Tripo AI API.')
bul(doc, 'FE-2: Render 3D furniture models directly within the user\'s physical space utilizing Google model-viewer for spatial WebAR.')

h3(doc, 'Module 7: Fraud Prevention')
bul(doc, 'FE-1: Execute zero-shot CLIP classification on uploaded images to guarantee they match the declared listing category.')
bul(doc, 'FE-2: Identify and block duplicate or stolen photos across the platform utilizing dHash perceptual image hashing.')
bul(doc, 'FE-3: Scan listing descriptions using NLP algorithms to actively flag scam-baiting terminology before a listing is made public.')

h3(doc, 'Module 8: Support & Notifications')
bul(doc, 'FE-1: Manage real-time system-wide alerts and visual unread count badges for all user activities.')
bul(doc, 'FE-2: Provide a dedicated support ticket dashboard allowing users to report issues and administrators to track resolution status.')

h3(doc, 'Module 9: Messaging & Inbox')
bul(doc, 'FE-1: Facilitate secure, real-time WebSocket-based messaging between prospective buyers and sellers.')
bul(doc, 'FE-2: Organize conversations within an inbox view grouped by specific listing to streamline negotiations.')

h3(doc, 'Module 10: Favorites & User Dashboard')
bul(doc, 'FE-1: Allow users to bookmark and save specific active listings into a dedicated favorites list for later viewing.')
bul(doc, 'FE-2: Provide a seller dashboard that tracks live status counts of the user\'s own active, pending, hidden, and sold advertisements.')

h3(doc, 'Module 11: AI Chatbot Assistant')
bul(doc, 'FE-1: Provide an intelligent customer support and market guide powered by the Groq llama-3.1-8b-instant LLM.')
bul(doc, 'FE-2: Stream real-time conversational responses to the client using Server-Sent Events (SSE).')
bul(doc, 'FE-3: Inject live OLX CSV market context directly into the prompt to ensure the chatbot provides accurate platform queries.')

doc.save(OUTPUT)
print(f'Chapter 2 done. Saved: {OUTPUT}')
