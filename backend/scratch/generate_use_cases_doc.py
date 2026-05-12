import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_use_cases_doc():
    doc = Document()

    # Title
    title = doc.add_heading('EZSell - Detailed Use Case Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    use_cases = [
        {
            "id": "UC-01",
            "name": "Post a Product for Sale",
            "actor": "Registered Seller",
            "goal": "Allow users to list their pre-owned product with AI-extracted specs and market-based price suggestions.",
            "pre": "User must be registered and logged in. Seller account verified.",
            "post": "Product listed and undergoing fraud detection.",
            "flow": [
                "1. Seller clicks 'Post New Product'.",
                "2. Seller enters product title (e.g., 'Samsung S24 Ultra').",
                "3. System uses LLM to validate title and extract specs (RAM, Storage, Color).",
                "4. System scrapes real-time market data from GSMArena and local Pakistani sites.",
                "5. System suggests a price range based on condition selections.",
                "6. Seller uploads images and confirms details.",
                "7. System runs automated fraud detection (CLIP/dHash).",
                "8. Product is listed and visible to buyers."
            ],
            "alt": "3a. If title is vague (e.g., 'Phone'), LLM provides specific improvement hints.",
            "exc": "Image format unsupported -> show error; Incomplete fields -> prompt user.",
            "pri": "High"
        },
        {
            "id": "UC-02",
            "name": "Search and Purchase a Product",
            "actor": "Registered Buyer",
            "goal": "Enable buyers to find specific products using advanced filters and initiate purchase.",
            "pre": "Internet connection active. User logged in.",
            "post": "Product marked 'In Negotiation' or 'Sold'.",
            "flow": [
                "1. Buyer enters keywords or uses category filters (Mobile, Laptop, Furniture).",
                "2. System applies location-based and category-specific filters.",
                "3. Buyer views results with thumbnails and engagement scores.",
                "4. Buyer selects a product to view detailed specs and AR preview (if furniture).",
                "5. Buyer clicks 'Contact Seller' to initiate chat.",
                "6. Parties negotiate price and delivery terms.",
                "7. Buyer confirms purchase; product status updated."
            ],
            "alt": "4a. AR try-on available for furniture -> buyer launches AR preview.",
            "exc": "No products found -> show alternative recommendations.",
            "pri": "High"
        },
        {
            "id": "UC-03",
            "name": "Detect Fraudulent Listing",
            "actor": "System (Fraud Detection), Admin",
            "goal": "Automatically detect and flag fraudulent or duplicate listings.",
            "pre": "Listing submitted by seller.",
            "post": "Listing approved, flagged for review, or rejected.",
            "flow": [
                "1. Listing data sent to Fraud Detection service.",
                "2. System analyzes text for suspicious keywords (e.g., 'free', 'advance payment').",
                "3. System runs CLIP to detect stock images or misleading visuals.",
                "4. System runs dHash to identify duplicate listings across the platform.",
                "5. If score > threshold, listing is approved automatically.",
                "6. If suspicious, listing is flagged for Admin manual review."
            ],
            "alt": "6a. Admin reviews flagged post and takes final action (Ban/Approve).",
            "exc": "ML Model timeout -> default to manual moderation.",
            "pri": "Medium-High"
        },
        {
            "id": "UC-04",
            "name": "Visualize Furniture in AR",
            "actor": "Registered Buyer",
            "goal": "Allow buyers to see how a furniture item fits in their actual physical environment.",
            "pre": "Device supports WebXR. Product has a 3D model.",
            "post": "Buyer gains confidence in fit and style.",
            "flow": [
                "1. Buyer clicks 'View in AR' on a furniture product page.",
                "2. System launches AR viewer and requests camera access.",
                "3. Buyer scans floor/walls to establish grounding.",
                "4. 3D model appears anchored to the surface.",
                "5. Buyer moves, rotates, and snaps the model to wall axes.",
                "6. System provides haptic feedback for alignment.",
                "7. Buyer takes screenshot or proceeds to contact seller."
            ],
            "alt": "4a. No 3D model exists -> System prompts to generate one from images.",
            "exc": "AR hardware incompatible -> show 3D web viewer fallback.",
            "pri": "High"
        },
        {
            "id": "UC-05",
            "name": "Secure In-App Messaging",
            "actor": "Buyer, Seller",
            "goal": "Facilitate real-time negotiation and communication between parties.",
            "pre": "Buyer initiates contact from listing page.",
            "post": "Chat history preserved for safety and reference.",
            "flow": [
                "1. Buyer sends a message to the seller.",
                "2. System delivers real-time notification to the seller.",
                "3. Seller opens Chat Window and responds.",
                "4. Parties exchange details, images, and price offers.",
                "5. System monitors for prohibited keywords (fraud prevention).",
                "6. Negotiation concludes with an agreement or cancellation."
            ],
            "alt": "5a. Automated chatbot handles common inquiries if seller is offline.",
            "exc": "Socket connection lost -> show retry indicator.",
            "pri": "High"
        },
        {
            "id": "UC-06",
            "name": "Market Price Suggestion Service",
            "actor": "System (Pricing Engine)",
            "goal": "Ground price suggestions in actual market data to prevent over/underpricing.",
            "pre": "Product title and category provided.",
            "post": "Fair market value range generated.",
            "flow": [
                "1. System triggers scraping of GSMArena (mobiles) or local tech sites.",
                "2. System gathers multiple price snippets from search results.",
                "3. System applies IQR (Interquartile Range) filtering to remove outliers.",
                "4. LLM analyzes filtered data to determine 'New' and 'Used' market prices.",
                "5. System factors in user-selected condition and specs.",
                "6. Final suggested range displayed to seller."
            ],
            "alt": "3a. No data found -> default to LLM general knowledge.",
            "exc": "Rate limit hit -> use cached price data.",
            "pri": "High"
        },
        {
            "id": "UC-07",
            "name": "Personalized Feed Generation",
            "actor": "System (Recommendation Service)",
            "goal": "Increase engagement by showing relevant products to users based on behavior.",
            "pre": "User has browsing/engagement history.",
            "post": "Personalized 'For You' section updated.",
            "flow": [
                "1. System tracks user clicks, views, and favorites.",
                "2. Engagement scores are calculated for categories and brands.",
                "3. Recommendation engine fetches listings matching high-score attributes.",
                "4. Feed is sorted by relevance and recency.",
                "5. Feed is served to the homepage dashboard."
            ],
            "alt": "3a. New user -> show trending/popular products.",
            "exc": "Profile data missing -> show general top listings.",
            "pri": "Medium"
        },
        {
            "id": "UC-08: Automatic 3D Model Generation",
            "id_num": "UC-08",
            "name": "Automated 3D Model Generation",
            "actor": "Registered Seller",
            "goal": "Transform 2D images into interactive 3D models for furniture listings.",
            "pre": "Seller uploads high-quality product images.",
            "post": "3D model (.glb) generated and attached to listing.",
            "flow": [
                "1. Seller selects 'Generate 3D' option.",
                "2. Images are sent to Tripo AI API via backend service.",
                "3. System polls API for generation progress.",
                "4. System downloads and optimizes the resulting mesh.",
                "5. Model is stored in cloud storage and linked to listing.",
                "6. Seller reviews model in 3D viewer."
            ],
            "alt": "4a. Generation quality low -> prompt for better lighting/angles.",
            "exc": "API credits exhausted -> notify user of delay.",
            "pri": "Medium-High"
        }
    ]

    for uc in use_cases:
        # Table A-2 Format
        doc.add_heading(f"Use Case {uc.get('id_num', uc['id'].split('-')[1])}: {uc['name']}", level=2)
        
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        data = [
            ("Field", "Description"),
            ("Use Case ID", uc.get('id_num', uc['id'])),
            ("Use Case Name", uc['name']),
            ("Primary Actor", uc['actor']),
            ("Goal", uc['goal']),
            ("Preconditions", uc['pre']),
            ("Postconditions", uc['post']),
            ("Main Flow", "\n".join(uc['flow'])),
            ("Alternate Flows", uc['alt']),
            ("Exceptions", uc['exc']),
            ("Priority", uc['pri'])
        ]
        
        for field, desc in data:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = desc
            # Make first column bold
            row_cells[0].paragraphs[0].runs[0].bold = True
            
        doc.add_paragraph() # Spacer

    doc.save('EZSell_Use_Cases.docx')
    print("EZSell_Use_Cases.docx created successfully.")

if __name__ == "__main__":
    create_use_cases_doc()
