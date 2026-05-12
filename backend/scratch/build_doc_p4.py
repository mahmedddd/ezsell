# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'c:\Users\ahmed\ezsell\EZSell_FYP_Final_Report.docx'
doc = Document(OUTPUT)

def cell_shade(cell, hex_col):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_col); tcPr.append(shd)

def h1(doc,t): return doc.add_heading(t,level=1)
def h2(doc,t): return doc.add_heading(t,level=2)
def h3(doc,t): return doc.add_heading(t,level=3)
def para(doc,t, style='Normal'): p=doc.add_paragraph(style=style); p.add_run(t); return p
def bpara(doc,t): p=doc.add_paragraph(); p.add_run(t).bold=True; return p
def bul(doc,t):
    p=doc.add_paragraph(style='List Paragraph')
    p.add_run(t); p.paragraph_format.left_indent=Inches(0.25); return p

def simple_table(doc, headers, rows, hcol='1F3864'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; cell_shade(c,hcol)
        for r in c.paragraphs[0].runs:
            r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]
        for ci,val in enumerate(row):
            tr.cells[ci].text=val
            if ri%2==0: cell_shade(tr.cells[ci],'DCE6F1')
            for r in tr.cells[ci].paragraphs[0].runs: r.font.size=Pt(9)
    return t

def placeholder(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f'>> [IMAGE PLACEHOLDER: {text}] <<')
    r.bold = True
    r.font.color.rgb = RGBColor(255,0,0)
    p.alignment = 1 # center

# ══════════════════════════════════════════════════════
# CHAPTER 4 – DESIGN AND ARCHITECTURE
# ══════════════════════════════════════════════════════
doc.add_page_break()
h1(doc,'4. Design and Architecture')
para(doc,
 'This chapter details the architectural and design blueprint of the EzSell system. '
 'It establishes the structural design paradigms, visualizes the operational flow through '
 'UML models, defines the data persistence schemas via a comprehensive data dictionary, '
 'and outlines the human-computer interface design.')

h2(doc,'4.1 Architectural Design')
para(doc,
 'EzSell follows a modern Multi-Tiered Client-Server Architecture, structurally '
 'divided into a Presentation Tier (React SPA), an Application Logic Tier (FastAPI), '
 'and a Data/AI Tier (PostgreSQL, ML models, External APIs).')
para(doc,
 '1. Presentation Layer (Frontend): Developed using React, Vite, and Tailwind CSS, '
 'this layer is responsible for rendering the UI, managing client-side routing, '
 'and handling 3D WebAR contexts using Google model-viewer.')
para(doc,
 '2. Logic Layer (Backend): Driven by FastAPI, this layer exposes RESTful endpoints, '
 'manages JWT authentication, handles form data parsing, and orchestrates calls to '
 'specialized service modules (e.g., llm_pricing_service, fraud_protection, image_to_3d).')
para(doc,
 '3. Data & Intelligence Layer: This encapsulates the PostgreSQL relational database '
 '(interfaced via SQLAlchemy ORM), the pre-trained ML ensemble (.pkl files), and '
 'external cognitive APIs (Groq LLM for language, Tripo AI for 3D generation).')

placeholder(doc, 'Box and Line Architecture Diagram showing Frontend, Backend, Database, ML Pipeline, and External APIs')

h2(doc,'4.2 Design Models')
para(doc, 'The system design is further elaborated using Object-Oriented Development design models.')

h3(doc,'Activity Diagram: Listing Creation & AI Validation')
para(doc,
 'This activity diagram maps the workflow of a user creating a new listing. It '
 'illustrates the initial input validation, the concurrent execution of the CLIP '
 'image-category validation and dHash duplicate checks, followed by the AI Price '
 'Prediction generation. If fraud checks fail, the activity terminates with a '
 'rejection; otherwise, it proceeds to the admin approval queue.')
placeholder(doc, 'Activity Diagram: Listing Creation Flow')

h3(doc,'Class Diagram: SQLAlchemy ORM Schema')
para(doc,
 'The class diagram represents the core entities in the SQLAlchemy Object Relational '
 'Mapping (ORM). It highlights the polymorphic relationships between the base `Listing` '
 'entity and the specialized categorical entities (`MobilePhone`, `Laptop`, `Furniture`), '
 'alongside associations with the `User`, `Message`, and `UserActivity` classes.')
placeholder(doc, 'Class Diagram: ORM Entities')

h3(doc,'Sequence Diagram: AR Try-On Generation (Tripo AI)')
para(doc,
 'This sequence diagram traces the asynchronous flow of AR model generation. It '
 'depicts the client requesting a 3D model, the backend uploading the image to '
 'Cloudinary, forwarding the URL to the Tripo AI API, and initiating a polling '
 'mechanism until the GLB asset is ready, compressed via Draco, and returned to the client.')
placeholder(doc, 'Sequence Diagram: Tripo AI API Interaction')

h3(doc,'State Transition Diagram: Listing Lifecycle')
para(doc,
 'This diagram visualizes the states of a product listing. The primary states are: '
 'Draft (client-side), Pending Review, Active, Hidden, Sold, and Rejected. State '
 'transitions are triggered by user actions (e.g., mark as sold) or admin actions '
 '(e.g., approve/reject).')
placeholder(doc, 'State Transition Diagram: Listing Status')

h2(doc,'4.3 Data Design')
para(doc,
 'The application utilizes a robust relational database schema designed to support '
 'rapid querying, geospatial filtering, and polymorphic category inheritance. The '
 'data is managed using PostgreSQL in production.')

h3(doc,'4.3.1 Data Dictionary')
para(doc, 'The following table lists the major data entities, their types, and descriptions within the EzSell database schema.')

para(doc, 'Table 4: Core Entities Data Dictionary', style='Caption')
simple_table(doc,
    ['Entity Name', 'Field', 'Data Type', 'Description'],
    [
        ['User', 'id', 'Integer (PK)', 'Unique primary key for the user.'],
        ['User', 'email', 'String', 'Verified email address, unique identifier.'],
        ['User', 'password_hash', 'String', 'Bcrypt hashed password string.'],
        ['User', 'is_verified', 'Boolean', 'True if OTP email verification is complete.'],
        ['Listing', 'id', 'Integer (PK)', 'Unique identifier for the product listing.'],
        ['Listing', 'owner_id', 'Integer (FK)', 'Foreign key referencing the User who created it.'],
        ['Listing', 'title', 'String', 'Groq-validated product title.'],
        ['Listing', 'price', 'Float', 'User-defined selling price in PKR.'],
        ['Listing', 'predicted_price', 'Float', 'AI-generated fair market price in PKR.'],
        ['Listing', 'status', 'Enum', 'Current state (pending, active, sold, hidden, rejected).'],
        ['MobilePhone', 'listing_id', 'Integer (FK, PK)', 'One-to-one link to base Listing.'],
        ['MobilePhone', 'ram_gb', 'Integer', 'System RAM extracted via NLP or user input.'],
        ['Furniture', 'material', 'String', 'Primary construction material (e.g., Wood, Metal).'],
        ['Message', 'id', 'Integer (PK)', 'Unique identifier for the direct message.'],
        ['Message', 'listing_id', 'Integer (FK)', 'The listing this conversation is about.'],
        ['Message', 'content', 'Text', 'The plaintext message body.'],
        ['UserActivity', 'activity_type', 'Enum', 'Type of interaction (view, favorite, message).'],
    ]
)

h2(doc,'4.4 Human Interface Design')
para(doc,
 'The human interface of EzSell focuses on reducing cognitive friction during complex '
 'tasks (pricing, 3D visualization) through clean, modern, light-themed typography '
 'and strategic micro-animations. The UI adheres to responsive web design principles.')

h3(doc,'4.4.1 Screen Images')
para(doc, 'The following sections present the core interfaces of the application from the user’s perspective.')
placeholder(doc, 'Screenshot: Home Page Feed with Search Filters')
placeholder(doc, 'Screenshot: Listing Creation Form with AI Pricing Widget')
placeholder(doc, 'Screenshot: Product Detail Page showing AR Try-On Button')
placeholder(doc, 'Screenshot: EzSell AI Chatbot Interface')

h3(doc,'4.4.2 Screen Objects and Actions')
para(doc, 'This section describes the interactive objects present on major screens and the user actions associated with them.')

para(doc, 'Table 5: Screen Objects and Actions', style='Caption')
simple_table(doc,
    ['Screen', 'Object', 'Action / Interaction'],
    [
        ['Home Feed', 'Category Pills', 'Clicking filters the masonry grid instantly to show only Mobiles, Laptops, or Furniture without page reload.'],
        ['Home Feed', 'Heart Icon (Card)', 'Toggles the listing in the user\'s favorites list; updates the server asynchronously.'],
        ['Listing Form', 'Predict Price Button', 'Triggers the backend ML pipeline. Displays a loading skeleton, then reveals the AI suggested price and confidence score.'],
        ['Listing Form', 'Image Dropzone', 'Supports drag-and-drop. Automatically scales thumbnails and triggers background CLIP validation.'],
        ['Detail Page', 'View in AR Button', 'Opens the full-screen Google model-viewer overlay. Prompts camera permissions and guides user to scan the floor.'],
        ['Detail Page', 'Message Seller Button', 'Opens a modal or navigates to the inbox, pre-filling the context with the selected listing\'s title.'],
        ['Chatbot', 'Floating Action Button', 'Expands the chat window. The interface supports Server-Sent Events (SSE) streaming for real-time text generation.'],
        ['Dashboard', 'Insights Tab', 'Renders Recharts gauges and donut charts visualizing the user\'s listing engagement metrics (views, clicks).'],
    ]
)

doc.save(OUTPUT)
print(f'Chapter 4 done. Saved: {OUTPUT}')
