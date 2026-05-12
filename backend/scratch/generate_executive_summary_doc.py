import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_executive_summary_doc():
    doc = Document()

    # Title
    title = doc.add_heading('Executive Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary_text = [
        "In the rapidly growing market for pre-owned goods, particularly within Pakistan, there is a persistent challenge in ensuring fair pricing, verifying product authenticity, and bridging the gap between digital listings and physical reality. Traditional marketplaces often suffer from price manipulation, fraudulent listings, and the \"uncertainty factor\" where buyers cannot visualize how items—especially furniture—will look in their actual environment. Relying on manual price research is time-consuming and often inaccurate, while static images fail to provide a comprehensive sense of scale and fit.",
        
        "To address these challenges and modernize the re-selling experience, EZSell has been developed. EZSell is an intelligent, AI-powered marketplace ecosystem designed to automate trust and enhance user engagement. It is a comprehensive platform that eliminates the guesswork for both sellers and buyers by integrating state-of-the-art machine learning and visualization technologies.",
        
        "EZSell provides a data-driven solution to the used-goods economy through three core pillars:",
        "1. Intelligent Pricing Engine: Utilizing Large Language Models (LLMs) and real-time market scraping from platforms like GSMArena and local marketplaces, EZSell provides sellers with statistically grounded price suggestions, preventing both overpricing and underpricing.",
        "2. Augmented Reality (AR) Integration: For furniture listings, the system allows buyers to virtually \"try on\" products in their own rooms. Using WebXR, users can place, rotate, and snap 3D models to their actual floor and walls, ensuring the product fits their space before purchase.",
        "3. Automated Integrity & Fraud Detection: The system employs advanced algorithms, including CLIP (Contrastive Language-Image Pretraining) and dHash (Difference Hashing), to analyze listing content and images. This ensures that fraudulent, duplicate, or misleading listings are flagged automatically before they reach the buyer.",
        
        "Technically, EZSell is built as a high-performance web application featuring a FastAPI (Python) backend and a React (Vite) frontend. It leverages Groq-powered Llama 3 models for semantic analysis and Tripo AI for automated 2D-to-3D model generation. The system operates on a real-time data pipeline that filters market outliers using IQR (Interquartile Range) analysis, ensuring that the insights provided to the user are both accurate and competitive."
    ]

    for paragraph in summary_text:
        doc.add_paragraph(paragraph)

    doc.save('EZSell_Executive_Summary.docx')
    print("EZSell_Executive_Summary.docx created successfully.")

if __name__ == "__main__":
    create_executive_summary_doc()
