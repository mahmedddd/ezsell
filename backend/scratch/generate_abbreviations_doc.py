import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_abbreviations_doc():
    doc = Document()

    # Title
    title = doc.add_heading('EZSell - List of Abbreviations', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    abbreviations = [
        ("LLM", "Large Language Model"),
        ("CLIP", "Contrastive Language-Image Pretraining"),
        ("dHash", "Difference Hashing"),
        ("IQR", "Interquartile Range"),
        ("WebXR", "Web Extended Reality"),
        ("GLB / GLTF", "Graphics Language Binary / Transmission Format"),
        ("PTA", "Pakistan Telecommunication Authority"),
        ("CPID", "Certificate Profile ID"),
        ("MDF", "Medium-Density Fibreboard"),
        ("PBR", "Physically Based Rendering"),
        ("JWT", "JSON Web Token"),
        ("CORS", "Cross-Origin Resource Sharing"),
        ("GZip", "GNU Zip Compression"),
        ("GSMArena", "Global System for Mobile Communications Association Arena")
    ]

    # Create Table
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    for abbr, full in abbreviations:
        row_cells = table.add_row().cells
        row_cells[0].text = abbr
        row_cells[1].text = full
        # Bold the abbreviation
        row_cells[0].paragraphs[0].runs[0].bold = True

    doc.save('EZSell_Abbreviations.docx')
    print("EZSell_Abbreviations.docx created successfully.")

if __name__ == "__main__":
    create_abbreviations_doc()
