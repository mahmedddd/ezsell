# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'c:\Users\ahmed\ezsell\EZSell_FYP_Final_Report.docx'
doc = Document(OUTPUT)

def cell_shade(cell, hex_col):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_col); tcPr.append(shd)

def h1(doc,t): return doc.add_heading(t,level=1)
def h2(doc,t): return doc.add_heading(t,level=2)
def h3(doc,t): return doc.add_heading(t,level=3)
def para(doc,t): p=doc.add_paragraph(); p.add_run(t); return p
def bul(doc,t):
    p=doc.add_paragraph(style='List Paragraph')
    p.add_run(t); p.paragraph_format.left_indent=Inches(0.25); return p

def fr_table(doc, identifier, title, req_desc):
    t = doc.add_table(rows=3, cols=2)
    t.style = 'Table Grid'
    # Row 0
    t.rows[0].cells[0].text = 'Identifier'
    t.rows[0].cells[1].text = identifier
    cell_shade(t.rows[0].cells[0], '1F3864')
    for r in t.rows[0].cells[0].paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    # Row 1
    t.rows[1].cells[0].text = 'Title'
    t.rows[1].cells[1].text = title
    cell_shade(t.rows[1].cells[0], '1F3864')
    for r in t.rows[1].cells[0].paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    # Row 2
    t.rows[2].cells[0].text = 'Requirement'
    t.rows[2].cells[1].text = req_desc
    cell_shade(t.rows[2].cells[0], '1F3864')
    for r in t.rows[2].cells[0].paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    return t

def caption(doc,t): return doc.add_paragraph(t,style='Caption')

# ══════════════════════════════════════════════════════
# CHAPTER 3 – REQUIREMENT ANALYSIS
# ══════════════════════════════════════════════════════
doc.add_page_break()
h1(doc,'3. Requirement Analysis')
para(doc,
 'This chapter outlines the detailed requirement analysis for the EzSell platform. '
 'It identifies the primary user classes, elaborates on the requirement elicitation '
 'techniques used, and exhaustively lists the functional and non-functional requirements '
 'vital for achieving the system’s goals. The requirements bridge the gap between '
 'psychological market needs (trust, pricing confidence, reduced cognitive load) and '
 'the practical implementation of AI and full-stack software architectures.')

h2(doc,'3.1 User classes and characteristics')
para(doc, 'The system caters to the following primary user classes, each with distinct psychological needs and operational characteristics:')

bul(doc, 'Buyers: Individuals seeking used electronics (mobiles, laptops) or furniture. '
         'Their primary concerns are trust (fear of scams), price fairness, and product condition. '
         'They require an intuitive UI, personalized discovery (reducing search fatigue), and AR try-on features '
         'to alleviate spatial anxiety when purchasing furniture.')
bul(doc, 'Sellers: Individuals looking to liquidate used assets quickly and profitably. '
         'They often experience pricing uncertainty and cognitive load when creating listings. '
         'They rely heavily on the AI Price Predictor to anchor their expectations and value a streamlined, '
         'category-specific listing flow.')
bul(doc, 'Administrators: Platform managers tasked with maintaining marketplace integrity. '
         'They require comprehensive analytical dashboards, automated fraud flags to prioritize moderation queues, '
         'and tools to resolve user disputes and support tickets efficiently.')

h2(doc,'3.2 Requirement Identifying Technique')
para(doc,
 'Use case modeling and event-response tables were selected as the primary techniques for '
 'requirement identification. Due to the highly interactive nature of the frontend '
 '(AR Try-On, AI Chatbot) and the complex, real-time backend pipelines (ML inference, '
 'CLIP fraud detection), these techniques effectively mapped user actions to systemic AI responses. '
 'Storyboarding was also utilized during the UI/UX design phase to map the cognitive flow '
 'of a seller traversing the AI pricing guidance step during listing creation.')

h2(doc,'3.3 Functional Requirements')
para(doc, 'The following functional requirements specify the software capabilities that must be implemented for the users to carry out the platform’s intended services.')

frs = [
    ('FR-1', 'User Registration and Authentication', 'The system shall allow users to register using an email and password or via Google OAuth 2.0. Standard registration requires mandatory phone number input and a 6-digit email OTP verification before listing creation is permitted. JWT Bearer tokens shall be used for session management.'),
    ('FR-2', 'AI-Powered Listing Creation', 'The system shall allow verified users to create listings with category-specific fields (e.g., RAM/Storage for Mobiles). The system shall automatically invoke the AI pricing engine to suggest a PKR market value and invoke the Groq LLM to validate the listing title.'),
    ('FR-3', 'Multi-Layer Fraud Prevention', 'The system shall process all uploaded images through a CLIP (ViT-B/32) model to ensure the image matches the declared category. Furthermore, a dHash perceptual algorithm shall detect and flag duplicate images across the database, while an NLP scanner flags scam-related keywords.'),
    ('FR-4', 'Machine Learning Price Prediction', 'The backend shall utilize a weighted ensemble model (XGBoost, LightGBM, Random Forest, Gradient Boosting) combined with Groq LLM live market scraping (filtered via IQR) to predict the fair market price of an item within an 85% accuracy threshold.'),
    ('FR-5', 'Augmented Reality (AR) Furniture Try-On', 'The system shall integrate with the Tripo AI API to procedurally generate a 3D GLB model from furniture listing photos. The frontend shall render this model via WebAR (model-viewer) and allow iOS users to download a USDZ file for native AR QuickLook.'),
    ('FR-6', 'Semantic Search and Personalized Feed', 'The system shall track user behavior (views, favorites) to build a time-decayed interest profile. It shall use SentenceTransformer embeddings to calculate cosine similarity between the user profile and available listings, presenting a personalized "For You" feed.'),
    ('FR-7', 'Admin Analytical Dashboard & Moderation', 'The system shall provide an administrative interface to view platform metrics, manage users, and process pending listings. Admins must be able to approve or reject listings (with attached rejection reasons).'),
    ('FR-8', 'In-App Messaging', 'The system shall provide a real-time messaging interface allowing buyers to communicate with sellers directly regarding specific listings, featuring unread message badges and thread management.'),
    ('FR-9', 'AI Chatbot Assistant', 'The system shall provide a persistent Chatbot (EzSell Assistant) powered by Groq (llama-3.1-8b-instant). The chatbot shall assist users with platform navigation and inject live OLX CSV data context to answer pricing-related queries.'),
    ('FR-10', 'Notification Management', 'The system shall trigger internal notifications to users regarding their listing approval status, received messages, and support ticket updates, accessible via a navbar notification bell.')
]

for identifier, title, desc in frs:
    h3(doc, f'{identifier}: {title}')
    caption(doc, f'Table {identifier[-1] if len(identifier)==4 else identifier[-2:]} Description of {identifier}')
    fr_table(doc, identifier, title, desc)
    doc.add_paragraph() # spacing

h2(doc,'3.4 Non-Functional Requirements')

h3(doc,'3.4.1 Reliability')
bul(doc, 'REL-1: The ML Pricing Engine shall provide a fallback to cached OLX CSV datasets if the DuckDuckGo live scraping or Groq API encounters rate limits or downtime.')
bul(doc, 'REL-2: The system shall ensure 99.9% uptime for core marketplace browsing and listing retrieval, relying on robust PostgreSQL clustering in production.')

h3(doc,'3.4.2 Usability')
bul(doc, 'USE-1: The AI Price Predictor shall present its estimate alongside a clear "Confidence Score" (0-100%) and a psychological anchor (e.g., "Fair Price", "Great Deal") to reduce cognitive load.')
bul(doc, 'USE-2: The AR Try-On interface shall feature an intuitive coaching overlay directing the user to scan their room floor before placing the 3D model.')

h3(doc,'3.4.3 Performance')
bul(doc, 'PER-1: The SentenceTransformer semantic search and recommendation engine shall return the personalized feed within 800 milliseconds for up to 10,000 active listings.')
bul(doc, 'PER-2: High-resolution images shall be optimized via Cloudinary CDN, ensuring the First Contentful Paint (FCP) of listing detail pages occurs under 1.5 seconds on 4G networks.')
bul(doc, 'PER-3: 3D GLB models generated by Tripo AI shall be compressed using Draco mesh optimization to keep file sizes under 5MB for rapid WebAR loading.')

h3(doc,'3.4.4 Security')
bul(doc, 'SEC-1: All passwords must be hashed using bcrypt with a salt factor of 12 before database persistence.')
bul(doc, 'SEC-2: JWT authentication tokens shall have a maximum lifespan of 24 hours, and sensitive administrative endpoints must strictly enforce role-based access control (RBAC).')
bul(doc, 'SEC-3: Uploaded images must be sanitized and stripped of EXIF GPS metadata to protect seller location privacy before being served to buyers.')

h2(doc,'3.5 External Interface Requirements')

h3(doc,'3.5.1 User Interfaces Requirements')
para(doc, 'The web application features a responsive, mobile-first UI constructed with React, Tailwind CSS, and Shadcn UI. Key interfaces include:')
bul(doc, 'Marketplace Feed: A masonry or grid layout of listing cards with visual price tags, condition badges, and quick-add favorite buttons.')
bul(doc, 'Listing Detail Page: A media-heavy interface with a central image carousel, AR Try-On button, AI price validation badge, seller trust metrics, and an integrated messaging widget.')
bul(doc, 'Admin Dashboard: A data-dense, desktop-optimized interface featuring Recharts graphs, moderation tables, and quick-action approval mechanisms.')

h3(doc,'3.5.2 Software interfaces')
para(doc, 'The system interacts with several external software services via RESTful APIs:')
bul(doc, 'Groq API: For high-speed LLM inference (Llama 3 models) used in price validation and the AI chatbot.')
bul(doc, 'Tripo AI API: For converting 2D furniture images into 3D GLB files.')
bul(doc, 'Cloudinary API: For secure, transformed, and CDN-cached media delivery.')
bul(doc, 'Google OAuth 2.0 API: For federated user authentication.')

h3(doc,'3.5.3 Hardware interfaces')
para(doc, 'While EzSell is a software platform, it interfaces with the user\'s hardware via standard web APIs:')
bul(doc, 'Camera API: Utilized on mobile devices for capturing listing photos directly from the web browser.')
bul(doc, 'Device GPU & Sensors: Utilized by Google model-viewer for WebGL rendering of 3D models and accelerometer/gyroscope tracking during AR placement.')

h3(doc,'3.5.4 Communications interfaces')
para(doc, 'The platform utilizes the following communication protocols:')
bul(doc, 'HTTP/1.1 and HTTP/2: For standard REST API requests and static asset delivery over TLS (HTTPS).')
bul(doc, 'Server-Sent Events (SSE): Employed for unidirectional real-time text streaming from the FastAPI backend to the frontend for the AI Chatbot responses.')
bul(doc, 'SMTP: For dispatching automated OTP verification and password reset emails to users.')

doc.save(OUTPUT)
print(f'Chapter 3 done. Saved: {OUTPUT}')
