# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'c:\Users\ahmed\ezsell\EZSell_FYP_Final_Report.docx'
doc = Document(OUTPUT)

def cell_shade(cell, hex_col):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_col); tcPr.append(shd)

def h1(doc,t): return doc.add_heading(t,level=1)
def h2(doc,t): return doc.add_heading(t,level=2)
def h3(doc,t): return doc.add_heading(t,level=3)
def para(doc,t, style='Normal'): p=doc.add_paragraph(style=style); p.add_run(t); return p
def bpara(doc,t): p=doc.add_paragraph(); p.add_run(t).bold=True; return p
def bul(doc,t):
    p=doc.add_paragraph(style='List Paragraph')
    p.add_run(t); p.paragraph_format.left_indent=Inches(0.25); return p

def simple_table(doc, headers, rows, hcol='1F3864'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; cell_shade(c,hcol)
        for r in c.paragraphs[0].runs:
            r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]
        for ci,val in enumerate(row):
            tr.cells[ci].text=val
            if ri%2==0: cell_shade(tr.cells[ci],'DCE6F1')
            for r in tr.cells[ci].paragraphs[0].runs: r.font.size=Pt(9)
    return t

# ══════════════════════════════════════════════════════
# CHAPTER 6 – TESTING AND EVALUATION
# ══════════════════════════════════════════════════════
doc.add_page_break()
h1(doc,'6. Testing and Evaluation')
para(doc,
 'Rigorous testing was conducted across all modules to ensure the system is stable, '
 'secure, and performs accurately under real-world conditions. This included unit testing '
 'of isolated AI scripts, functional testing of UI flows, and complex integration '
 'testing across third-party APIs.')

h2(doc,'6.1 Unit Testing')
para(doc, 'Unit testing validates that individual components, especially the isolated ML and CV scripts, perform as expected.')
para(doc, 'Table 10: Unit Testing Results', style='Caption')
simple_table(doc,
    ['No.', 'Test case / Module', 'Attribute and value', 'Expected result', 'Result'],
    [
        ['1', 'CLIP Zero-Shot Image Classifier', 'Image: laptop.jpg, Labels: [Laptop, Furniture, Mobile]', 'Returns highest probability for "Laptop".', 'Pass'],
        ['2', 'dHash Image Deduplication', 'Two identical images compressed differently', 'Hamming distance < 5 (Flags as duplicate).', 'Pass'],
        ['3', 'Password Hashing (bcrypt)', 'Password string input', 'Outputs 60-char hash; verifying identical password returns True.', 'Pass'],
        ['4', 'IQR Outlier Filter', 'List of prices: [1000, 1050, 1100, 95000]', 'Returns [1000, 1050, 1100], removing 95000.', 'Pass'],
    ]
)

h2(doc,'6.2 Functional Testing')
para(doc, 'Functional testing ensures that the system interfaces meet the functional specifications detailed in the requirements phase.')
para(doc, 'Table 11: Functional Testing Results', style='Caption')
simple_table(doc,
    ['No.', 'Test case / Module', 'Attribute and value', 'Expected result', 'Actual result', 'Result'],
    [
        ['1', 'User Login Flow', 'Email: test@ezsell.com\nPass: valid123', 'Authenticates, returns JWT, redirects to Home Feed.', 'Redirected to Home with active session state.', 'Pass'],
        ['2', 'Category Form Toggle', 'Select "Furniture" from dropdown', 'UI dynamically renders Material and Subtype fields.', 'Furniture fields rendered correctly.', 'Pass'],
        ['3', 'Create Listing (No OTP)', 'Submit listing with unverified account', 'HTTP 403 Forbidden with "Verify Email" error.', '403 error displayed on UI alert.', 'Pass'],
        ['4', 'AR Viewer Button', 'Click "View in AR" on mobile device', 'model-viewer requests camera overlay.', 'Camera overlay opened successfully.', 'Pass'],
    ]
)

h2(doc,'6.3 Business Rules Testing')
para(doc, 'Decision table based testing is used to validate the complex business logic of the Multi-Layer Fraud Protection Pipeline.')
para(doc, 'Table 12: Fraud Pipeline Decision Matrix', style='Caption')
simple_table(doc,
    ['Rule', 'Is Verified?', 'Duplicate Hash?', 'CLIP Match?', 'Scam Text?', 'System Action', 'Result'],
    [
        ['1', 'No', 'N/A', 'N/A', 'N/A', 'Hard Reject (403)', 'Pass'],
        ['2', 'Yes', 'Yes', 'N/A', 'N/A', 'Reject (Duplicate)', 'Pass'],
        ['3', 'Yes', 'No', 'No', 'N/A', 'Reject (Category Mismatch)', 'Pass'],
        ['4', 'Yes', 'No', 'Yes', 'Yes', 'Pending Review (Flagged)', 'Pass'],
        ['5', 'Yes', 'No', 'Yes', 'No', 'Approved (Active)', 'Pass'],
    ]
)

h2(doc,'6.4 Integration Testing')
para(doc, 'Integration testing validates the hand-offs between the FastAPI backend and external intelligence APIs.')
para(doc, 'Table 13: Integration Testing Results', style='Caption')
simple_table(doc,
    ['No.', 'Test case / Flow', 'Attribute and value', 'Expected result', 'Actual result', 'Result'],
    [
        ['1', 'LLM Pricing Pipeline', 'Submit Mobile Specs to backend', 'Backend scrapes OLX, queries Groq API, returns PKR price.', 'Returned accurate PKR price in < 3s.', 'Pass'],
        ['2', 'Tripo 3D Generation', 'Upload furniture image to backend', 'Image to Cloudinary -> URL to Tripo -> Polling -> GLB save.', 'GLB URL returned and saved in DB.', 'Pass'],
        ['3', 'Chatbot SSE Stream', 'Send message to /chat/stream', 'Frontend receives chunked Server-Sent Events from Groq.', 'Text streams dynamically in UI.', 'Pass'],
    ]
)

# ══════════════════════════════════════════════════════
# CHAPTER 7 – CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════
doc.add_page_break()
h1(doc,'7. Conclusion and Future Work')

h2(doc,'7.1 Conclusion')
para(doc,
 'The EzSell project successfully conceptualizes, designs, and implements a '
 'next-generation, AI-powered second-hand marketplace tailored for the Pakistani '
 'economy. By bridging the gap between raw classifieds and structured e-commerce, '
 'the platform fundamentally resolves the "trust deficit" and "pricing anxiety" '
 'endemic to the informal sector.')
para(doc,
 'Through the successful integration of a dual-track ML and Groq LLM pricing engine, '
 'sellers are empowered with data-grounded PKR market intelligence, eliminating '
 'decision fatigue. The Multi-Layer Fraud Pipeline, leveraging zero-shot CLIP '
 'classification and perceptual hashing, proactively sterilizes the marketplace '
 'of scams before they reach buyers. Furthermore, the pioneering inclusion of '
 'Tripo AI-driven 3D generation and WebAR visualization drastically reduces spatial '
 'anxiety for furniture buyers, offering an immersive try-before-you-buy experience.')
para(doc,
 'Operating on a robust FastAPI and React architecture, the project proves that '
 'advanced cognitive services can be seamlessly woven into consumer web applications '
 'to create secure, structured, and highly intelligent digital ecosystems.')

h2(doc,'7.2 Future Work')
para(doc, 'While EzSell 1.0.0 delivers a comprehensive marketplace, the architecture is designed for extensive future scaling:')
bul(doc, 'In-App Escrow Payments: Integrating payment gateways (e.g., Stripe, JazzCash) to allow secure, platform-held funds until the buyer verifies the item in person.')
bul(doc, 'Native Mobile Applications: Porting the React frontend to React Native to deliver dedicated iOS and Android applications, utilizing native camera APIs for advanced LiDAR 3D scanning.')
bul(doc, 'Expanded Category Intelligence: Training ML ensemble models for new high-value categories, including Vehicles and Real Estate.')
bul(doc, 'Live Video Commerce: Implementing WebRTC to allow sellers to host live-streamed "virtual garage sales" directly within the platform.')

# ══════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════
doc.add_page_break()
h1(doc,'8. References')
refs = [
 'Radford, A., et al. (2021). "Learning Transferable Visual Models From Natural Language Supervision." (CLIP Architecture). arXiv:2103.00020.',
 'Reimers, N., & Gurevych, I. (2019). "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks." Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing. (SentenceTransformers).',
 'Chen, T., & Guestrin, C. (2016). "XGBoost: A Scalable Tree Boosting System." Proceedings of the 22nd ACM SIGKDD International Conference. (XGBoost logic).',
 'Ke, G., et al. (2017). "LightGBM: A Highly Efficient Gradient Boosting Decision Tree." Advances in Neural Information Processing Systems. (LightGBM application).',
 'FastAPI Documentation. "FastAPI framework, high performance, easy to learn, fast to code." https://fastapi.tiangolo.com/',
 'Meta (2024). Llama 3 Model Architecture. Integrated via Groq LPU API Documentation. https://console.groq.com/docs',
 'Tripo AI API V2 Documentation. "Image-to-3D Generation Pipeline." https://platform.tripo3d.ai/docs/api',
]
for r in refs:
    p = doc.add_paragraph(style='List Paragraph')
    p.add_run(r)

doc.save(OUTPUT)
print(f'Chapters 6, 7, 8 done. Saved: {OUTPUT}')
