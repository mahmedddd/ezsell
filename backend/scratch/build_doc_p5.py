# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'c:\Users\ahmed\ezsell\EZSell_FYP_Final_Report.docx'
doc = Document(OUTPUT)

def cell_shade(cell, hex_col):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_col); tcPr.append(shd)

def h1(doc,t): return doc.add_heading(t,level=1)
def h2(doc,t): return doc.add_heading(t,level=2)
def h3(doc,t): return doc.add_heading(t,level=3)
def para(doc,t, style='Normal'): p=doc.add_paragraph(style=style); p.add_run(t); return p
def bpara(doc,t): p=doc.add_paragraph(); p.add_run(t).bold=True; return p
def bul(doc,t):
    p=doc.add_paragraph(style='List Paragraph')
    p.add_run(t); p.paragraph_format.left_indent=Inches(0.25); return p

def simple_table(doc, headers, rows, hcol='1F3864'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; cell_shade(c,hcol)
        for r in c.paragraphs[0].runs:
            r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]
        for ci,val in enumerate(row):
            tr.cells[ci].text=val
            if ri%2==0: cell_shade(tr.cells[ci],'DCE6F1')
            for r in tr.cells[ci].paragraphs[0].runs: r.font.size=Pt(9)
    return t

def algo_table(doc, title, inputs, outputs, steps):
    ncols = 3
    nrows = 3 + len(steps)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = 'Table Grid'
    t.rows[0].cells[0].merge(t.rows[0].cells[ncols-1])
    c = t.rows[0].cells[0]
    c.text = title
    cell_shade(c, '1F3864')
    for run in c.paragraphs[0].runs:
        run.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(10)
    t.rows[1].cells[0].merge(t.rows[1].cells[ncols-1])
    t.rows[1].cells[0].text = 'Input: ' + inputs
    t.rows[2].cells[0].merge(t.rows[2].cells[ncols-1])
    t.rows[2].cells[0].text = 'Output: ' + outputs
    for i, step in enumerate(steps):
        t.rows[3+i].cells[0].merge(t.rows[3+i].cells[ncols-1])
        t.rows[3+i].cells[0].text = f"Step {i+1}: {step}"
    return t

def placeholder(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f'>> [IMAGE PLACEHOLDER: {text}] <<')
    r.bold = True; r.font.color.rgb = RGBColor(255,0,0)
    p.alignment = 1

# ══════════════════════════════════════════════════════
# CHAPTER 5 – IMPLEMENTATION
# ══════════════════════════════════════════════════════
doc.add_page_break()
h1(doc,'5. Implementation')
para(doc,
 'This chapter explores the practical implementation details of the EzSell system. '
 'It outlines the core algorithms governing pricing and security, details the external '
 'APIs integrated into the platform, reviews the structural design of the user interface, '
 'and provides the deployment strategy for the production environment.')

h2(doc,'5.1 Algorithm')
para(doc,
 'The platform relies heavily on autonomous data-driven algorithms to ensure marketplace '
 'integrity and precision pricing. Pseudocode for the three most critical systemic '
 'algorithms is provided below.')

h3(doc, 'Algorithm 1: Dual-Track AI Price Prediction & IQR Filtering')
algo_table(doc,
 'Algorithm 1: AI_Price_Prediction_Pipeline',
 'L (Listing Attributes: Title, Category, Specs)',
 'P (Predicted Fair Market Price in PKR), C (Confidence Score 0-100%)',
 [
  'Initialize ML_Ensemble weights: w_xgb=0.35, w_lgb=0.35, w_rf=0.15, w_gb=0.15',
  'Load serialized pre-trained models (XGBoost, LightGBM, RandomForest, GradientBoosting)',
  'Base_Price = (w_xgb*XGB_pred(L)) + (w_lgb*LGB_pred(L)) + (w_rf*RF_pred(L)) + (w_gb*GB_pred(L))',
  'Search_Query = Construct_Query_From_Attributes(L)',
  'Scraped_Prices[] = Scrape_DuckDuckGo_OLX(Search_Query)',
  'If length(Scraped_Prices) < 5:',
  '    Return Base_Price, Confidence = 60%  # Low data fallback',
  'Q1, Q3 = Calculate_Quartiles(Scraped_Prices)',
  'IQR = Q3 - Q1',
  'Filtered_Prices[] = Remove prices outside [Q1 - 1.5*IQR, Q3 + 1.5*IQR]',
  'LLM_Market_Price = Average(Filtered_Prices)',
  'Final_Price = (Base_Price * 0.85) + (LLM_Market_Price * 0.15)',
  'Calculate Confidence based on variance between Base_Price and LLM_Market_Price',
  'Return Final_Price, Confidence'
 ])
para(doc, 'Table 6: Dual-Track Pricing Algorithm', style='Caption')

h3(doc, 'Algorithm 2: Multi-Layer Fraud Protection Pipeline')
algo_table(doc,
 'Algorithm 2: Fraud_Protection_Pipeline',
 'L (New Listing: Title, Desc, Price, Image_Uploads[])',
 'Status (Approved, Rejected, Pending_Review), Flags[]',
 [
  'Initialize Flags[] as empty list',
  'Check 1: If User.is_verified == False -> Return Rejected',
  'Check 2: L_Hash = MD5(L.Title + L.Desc + L.Price + L.Owner_ID)',
  'If L_Hash exists in DB -> Return Rejected (Copy-paste duplicate)',
  'Check 3: For each Img in Image_Uploads[]:',
  '    Prediction = CLIP_Model.Predict(Img, Candidate_Labels=[L.Category, "Other"])',
  '    If Prediction != L.Category -> Add "Category Mismatch" to Flags[]',
  'Check 4: For each Img in Image_Uploads[]:',
  '    Img_dHash = compute_perceptual_dhash(Img)',
  '    If Img_dHash exists in Global_Image_Hashes -> Add "Stolen Image" to Flags[]',
  'Check 5: If L.Desc contains Scam_Regex_Patterns -> Add "Scam Keywords" to Flags[]',
  'If len(Flags[]) > 0 -> Return Pending_Review, Flags[]',
  'Else -> Return Approved, []'
 ])
para(doc, 'Table 7: Fraud Protection Algorithm', style='Caption')

h3(doc, 'Algorithm 3: Hybrid Semantic Recommendation')
algo_table(doc,
 'Algorithm 3: Hybrid_Recommendation_Scoring',
 'U (User Activity History), L_Pool[] (Available Active Listings)',
 'Ranked_Listings[]',
 [
  'U_Interests[] = Extract categories and keywords from U.Activity_History',
  'Apply Time-Decay Weighting: Weight = Activity_Score * exp(-lambda * Days_Since_Activity)',
  'U_Vector = SentenceTransformer.Encode(U_Interests with Weights)',
  'For each listing L in L_Pool[]:',
  '    L_Vector = Retrieve pre-computed SentenceTransformer embedding for L',
  '    Similarity_Score = Cosine_Similarity(U_Vector, L_Vector)',
  '    Popularity_Boost = Log(L.view_count + 1) * 0.1',
  '    Final_Score = Similarity_Score + Popularity_Boost',
  'Sort L_Pool[] descending by Final_Score -> Ranked_Listings[]',
  'Return Ranked_Listings[]'
 ])
para(doc, 'Table 8: Recommendation Algorithm', style='Caption')

h2(doc,'5.2 External APIs/SDKs')
para(doc, 'EzSell integrates several third-party services to handle high-compute AI processing and secure media management.')
para(doc, 'Table 9: Details of APIs used in the project', style='Caption')
simple_table(doc,
    ['Name of API', 'Description', 'Purpose of Usage', 'Function / Endpoint'],
    [
        ['Groq API', 'Ultra-fast LPU inference engine for open-source LLMs.', 'Runs llama-3.3-70b for price validation and llama-3.1-8b for the chatbot.', 'chat.completions.create'],
        ['Tripo AI V2', '2D to 3D Generative AI service.', 'Generates GLB 3D models from furniture photos for AR try-on.', 'https://api.tripo3d.ai/v2/openapi/task'],
        ['Cloudinary', 'Cloud media CDN and transformation API.', 'Stores images/GLB files, serves them via CDN, and performs auto-scaling.', 'cloudinary.uploader.upload'],
        ['Google OAuth 2.0', 'Federated identity provider.', 'Allows users to sign up and log in securely via their Google accounts.', 'accounts.google.com/o/oauth2/v2/auth'],
        ['DuckDuckGo DDGS', 'Search scraping library.', 'Scrapes live OLX Pakistan links to retrieve real-time market prices.', 'ddgs.text()'],
    ]
)

h2(doc,'5.3 User Interface')
para(doc,
 'The user interface is constructed dynamically using React 18 and Tailwind CSS, adhering '
 'to a unified "Light Theme" design system to convey professionalism and trustworthiness. '
 'Below are descriptions of key interfaces across the sub-systems.')

h3(doc, '5.3.1 Client Web App: Home Feed')
para(doc, 'The Home Feed is the entry point, featuring a responsive masonry layout. It implements infinite scrolling and instantaneous client-side filtering via category pills.')
placeholder(doc, 'Screenshot: Client Web App Home Feed')

h3(doc, '5.3.2 Client Web App: AR Viewer')
para(doc, 'The AR Viewer utilizes Google model-viewer overlayed on the screen. It provides a 360-degree interactive canvas and a prominent "View in your space" AR button triggering native camera sessions.')
placeholder(doc, 'Screenshot: AR Viewer Interface')

h3(doc, '5.3.3 Admin Web App: Analytical Dashboard')
para(doc, 'The Admin dashboard relies on Recharts to render performance gauges and donut charts. It provides a tabular view of all pending listings with quick-action Approval/Rejection buttons and fraud flag indicators.')
placeholder(doc, 'Screenshot: Admin Dashboard')

h2(doc,'5.4 Deployment')
para(doc,
 'The deployment architecture is designed for scalability and high availability, '
 'separating the frontend delivery from the backend computational nodes.')
bul(doc, 'Backend API Hosting: The FastAPI application is containerized via Docker and deployed on a cloud virtual private server (e.g., AWS EC2 or DigitalOcean Droplet) managed by Gunicorn with Uvicorn workers for asynchronous throughput.')
bul(doc, 'Frontend Delivery: The React/Vite SPA is statically built and deployed to Vercel, utilizing their edge network for instantaneous global delivery and HTTPS provisioning.')
bul(doc, 'Database: A managed PostgreSQL instance (e.g., AWS RDS or Supabase) ensures data durability, automated backups, and transactional integrity.')
bul(doc, 'Media Storage: All static assets (images, GLB, USDZ files) are offloaded to Cloudinary, which provides on-the-fly transformations and aggressive 1-year cache headers for heavy 3D assets to minimize bandwidth costs.')

doc.save(OUTPUT)
print(f'Chapter 5 done. Saved: {OUTPUT}')
