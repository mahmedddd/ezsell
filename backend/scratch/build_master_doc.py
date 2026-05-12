# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from itertools import groupby

TEMPLATE = r'c:\Users\ahmed\ezsell\FYP_Final Report_Sring 2026_Template (Updated) (1).docx'
OUTPUT   = r'c:\Users\ahmed\ezsell\EZSell_FYP_Final_Report.docx'

doc = Document(TEMPLATE)

# TRUNCATE AT INTRO (Index 180)
body = doc.element.body
children = list(body)
para_count = 0; cut_idx = None
for i, child in enumerate(children):
    if child.tag.endswith('}p') or child.tag.endswith('}tbl') or child.tag.endswith('}sectPr'):
        if child.tag.endswith('}p'):
            if para_count == 180:
                cut_idx = i
                break
            para_count += 1

if cut_idx is not None:
    to_remove = list(children)[cut_idx:]
    secPrs = [el for el in to_remove if el.tag.endswith('}sectPr')]
    for el in to_remove: body.remove(el)
    if secPrs: body.append(secPrs[-1])

# HELPERS
def cell_shade(cell, hex_col):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
    shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_col); tcPr.append(shd)

def h1(doc,t): return doc.add_heading(t,level=1)
def h2(doc,t): return doc.add_heading(t,level=2)
def h3(doc,t): return doc.add_heading(t,level=3)
def para(doc,t, style='Normal'): 
    p=doc.add_paragraph(style=style)
    p.add_run(t)
    p.paragraph_format.space_after = Pt(10)
    return p

def bul(doc,t):
    p = doc.add_paragraph(style='List Paragraph')
    p.add_run('•  ' + t)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    return p

def simple_table(doc, headers, rows, hcol='1F3864'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; cell_shade(c,hcol)
        for r in c.paragraphs[0].runs:
            r.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        tr=t.rows[ri+1]
        for ci,val in enumerate(row):
            tr.cells[ci].text=val
            if ri%2==0: cell_shade(tr.cells[ci],'DCE6F1')
            for r in tr.cells[ci].paragraphs[0].runs: r.font.size=Pt(9)
    doc.add_paragraph() # Add spacing after table
    return t

def algo_table(doc, title, inputs, outputs, steps):
    """Pseudocode-style algorithm table matching exact image format."""
    t = doc.add_table(rows=4, cols=1)
    t.style = 'Table Grid'
    
    # Title row
    c0 = t.rows[0].cells[0]
    c0.text = title
    cell_shade(c0, '404040')
    for run in c0.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        
    # Input row
    c1 = t.rows[1].cells[0]
    c1.text = 'Input: ' + inputs
    for run in c1.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
        
    # Output row
    c2 = t.rows[2].cells[0]
    c2.text = 'Output: ' + outputs
    for run in c2.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(10)
        
    # Steps row (one big cell)
    c3 = t.rows[3].cells[0]
    # Join all steps with newlines
    c3.text = '\n'.join(steps)
    for run in c3.paragraphs[0].runs:
        run.font.size = Pt(10)
        
    return t

def placeholder(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f'>> [IMAGE PLACEHOLDER: {text}] <<')
    r.bold = True; r.font.color.rgb = RGBColor(255,0,0)
    p.alignment = 1
    p.paragraph_format.space_after = Pt(24)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 – INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, 'Introduction')
para(doc, 'This chapter provides an overview of the EzSell project, including its vision, a comparative analysis of existing marketplace systems, key deliverables, system constraints, technology stack, and relevance to the academic curriculum. EzSell is a full-stack, AI-powered second-hand marketplace platform tailored for the Pakistani market, enabling intelligent buying and selling of used mobile phones, laptops, and furniture through a secure, visually rich web application.')

h2(doc, 'Vision Statement')
para(doc, 'For individual buyers and sellers in Pakistan who face challenges in pricing second-hand electronics and furniture accurately and safely, EzSell is an AI-powered web marketplace that provides real-time price intelligence, automated fraud prevention, AR visualization, and a personalized buying experience. Unlike existing platforms such as OLX and Facebook Marketplace, which lack AI-assisted pricing, structured fraud detection, and immersive product visualization, EzSell leverages a multi-model ML ensemble, a Groq LLM layer, CLIP-based image validation, SentenceTransformer semantic recommendations, and Tripo AI 3D generation to deliver a trustworthy, data-driven, and visually superior marketplace for the pre-owned goods economy of Pakistan.')

h2(doc, 'Related System Analysis / Literature Review')
para(doc, 'The following table presents a comparative analysis of leading marketplace platforms against the proposed EzSell solution, identifying their weaknesses and the specific innovations EzSell introduces.')

para(doc, 'Table 1   Related System Analysis with Proposed Project Solution', style='Caption')
simple_table(doc,
    ['Application Name', 'Weakness', 'Proposed Project Solution'],
    [
        ['OLX Pakistan', 'No AI-assisted pricing; prices are seller-determined leading to widespread overpricing. High incidence of scam listings with no automated image verification.', 'EzSell integrates a dual-track ML + LLM pricing pipeline (XGBoost, LightGBM, Groq) with live OLX market scraping. CLIP-based image validation eliminates fraudulent listings.'],
        ['Facebook Marketplace', 'Lacks structured pricing guidance. No category-specific fields for electronics. Buyers cannot visualize furniture in their actual space.', 'EzSell provides category-specific attribute fields, AI price suggestions, and a full AR Try-On viewer powered by Tripo AI and Google model-viewer for in-room furniture visualization.'],
        ['Daraz.pk', 'Primarily a new-goods B2C platform; does not support C2C used-goods listings. No secondhand ecosystem.', 'EzSell is purpose-built for the secondhand C2C market, supporting individual sellers with full listing lifecycle management and AI-assisted used-goods pricing.'],
    ]
)

h2(doc, 'Project Deliverables')
para(doc, 'The following deliverables constitute the complete EzSell 1.0.0 production release:')
dels = [
    'D-1: Fully functional FastAPI backend with RESTful APIs for all marketplace operations.',
    'D-2: React 18 + Vite + TypeScript frontend SPA with Tailwind CSS and Shadcn UI.',
    'D-3: AI-powered price prediction engine combining a trained ML ensemble and Groq LLM validation.',
    'D-4: Automated fraud prevention pipeline utilizing CLIP category validation, dHash duplicate detection, and scam keyword filtering.',
    'D-5: Semantic recommendation engine using SentenceTransformer embedding-based personalized feeds.',
    'D-6: AR/3D furniture try-on module featuring Tripo AI GLB generation and WebAR viewer integration.',
    'D-7: Admin analytical dashboard equipped with listing moderation, user management, and platform engagement metrics.',
    'D-8: Real-time in-app messaging system providing direct buyer-to-seller communication.',
    'D-9: AI chatbot assistant (EzSell Assistant) powered by Groq llama-3.1-8b-instant with live context injection.',
    'D-10: Cloudinary-integrated cloud media storage with optimized content delivery.'
]
for d in dels: bul(doc, d)

h2(doc, 'System Limitations / Constraints')
limits = [
    'L-1: EzSell is a web-only platform at launch; no dedicated iOS or Android native application is provided, though the frontend is fully mobile-responsive.',
    'L-2: Only three product categories are supported in v1.0.0 — Mobiles, Laptops, and Furniture.',
    'L-3: Web scraping of OLX Pakistan via DuckDuckGo is subject to external rate limits; the system falls back to cached CSV data if rate-limited.',
    'L-4: The platform does not facilitate in-app payment processing; all financial transactions are completed offline between the buyer and seller.',
    'L-5: AR 3D generation via the Tripo AI API requires significant compute time, taking approximately 30–90 seconds per generation depending on third-party API availability.'
]
for l in limits: bul(doc, l)

h2(doc, 'Tools and Technologies')
para(doc, 'Table 2 below lists all hardware and software tools, grouped by category, used in the implementation of EzSell.')
para(doc, 'Table 2   Tools and Technologies for EzSell', style='Caption')

tools_data = [
    ('Programming Language', 'Python', '3.11+', 'Backend logic, ML pipeline, AI services'),
    ('Programming Language', 'TypeScript', '5.8', 'Type-safe frontend development'),
    ('Programming Language', 'JavaScript', 'ES2023', 'Frontend runtime interactions'),
    ('Frontend', 'React', '18.3', 'Component-based SPA architecture'),
    ('Frontend', 'Vite', '5.4', 'Ultra-fast HMR dev server and bundler'),
    ('Frontend', 'Shadcn UI', 'Latest', 'Accessible headless UI primitives'),
    ('Frontend', 'Tailwind CSS', '3.4', 'Utility-first styling'),
    ('3D / AR', 'Google model-viewer', '4.1', 'WebAR GLB rendering'),
    ('3D / AR', 'Three.js / React-Three-Fiber', '0.182 / 8.18', 'Procedural 3D scenes and WebGL'),
    ('Backend', 'FastAPI', '0.115+', 'Async Python REST API'),
    ('Backend', 'Uvicorn', 'Latest', 'Production ASGI server'),
    ('Backend', 'SQLAlchemy / Pydantic', '2.x', 'ORM and schema validation'),
    ('Database', 'PostgreSQL / SQLite', '15 / 3.x', 'Production and Dev data storage'),
    ('AI / ML', 'XGBoost / LightGBM', '2.x', 'Primary gradient boosted price prediction (35% weight each)'),
    ('AI / ML', 'Scikit-learn', '1.x', 'Random Forest (15%) and GradientBoosting (15%) price prediction'),
    ('AI / ML', 'SentenceTransformers', '3.x', 'Semantic embedding model'),
    ('AI / ML', 'Groq API', 'Latest', 'llama-3.3-70b (pricing); llama-3.1-8b (chatbot)'),
    ('AI / ML', 'OpenAI CLIP', 'Latest', 'Zero-shot image classification'),
    ('AI / ML', 'Tripo AI API', 'V2', 'Image-to-3D GLB generation'),
    ('Infrastructure', 'AWS EC2', 'c7i-flex.large', 'Cloud compute server for backend API'),
    ('Infrastructure', 'Cloudinary', 'Latest', 'Cloud media CDN and transformation')
]

t_tools = doc.add_table(rows=1, cols=3)
t_tools.style = 'Table Grid'
hr = t_tools.rows[0]
for i,h in enumerate(['Tool / Technology', 'Version', 'Rationale']):
    hr.cells[i].text = h
    cell_shade(hr.cells[i], '1F3864')
    for r in hr.cells[i].paragraphs[0].runs: r.bold=True; r.font.color.rgb=RGBColor(255,255,255)

for category, items in groupby(tools_data, key=lambda x: x[0]):
    row = t_tools.add_row()
    row.cells[0].merge(row.cells[2])
    row.cells[0].text = category
    cell_shade(row.cells[0], 'E2E8F0')
    for r in row.cells[0].paragraphs[0].runs: r.bold=True
    
    for item in items:
        row = t_tools.add_row()
        row.cells[0].text = item[1]
        row.cells[1].text = item[2]
        row.cells[2].text = item[3]
doc.add_paragraph()

h2(doc, 'Relevance to Course Modules')
para(doc, 'EzSell integrates extensive knowledge from the core and elective modules of the BS Artificial Intelligence (Fall 2022) curriculum at Air University. The system directly applies theoretical concepts into a production-ready application:')
crmods = [
    'AI332 Machine Learning: Directly applied in the dual-track pricing engine, utilizing ensemble methods (XGBoost, LightGBM, Random Forest, Gradient Boosting) to train predictive models on scraped second-hand market data.',
    'AI320 Natural Language Processing: Implemented for semantic analysis within the multi-layer fraud pipeline (scam keyword detection) and the Groq LLM-powered conversational chatbot assistant.',
    'AI303 Computer Vision: Utilized for the zero-shot image classification pipeline (OpenAI CLIP) to validate product categories and dHash perceptual hashing to identify duplicate or stolen listing photos.',
    'SE100 Software Engineering & CS112 Object Oriented Programming: Formed the foundation of the project’s SDLC, architectural design, and the object-oriented structure of the FastAPI backend and React frontend SPA.',
    'CS230 Database Systems: Applied in designing the normalized PostgreSQL relational database schema and mapping complex polymorphic relationships using the SQLAlchemy ORM.',
    'CS415 Information Security: Implemented to secure user data via bcrypt password hashing, JWT-based stateless session management, and robust API endpoint protection.',
    'CS214 Data Structures and Algorithms: Essential for developing the hybrid semantic recommendation engine, optimizing cosine similarity mathematical operations, and executing rapid in-memory vector lookups.',
    'AI301 Knowledge Representation and Reasoning: Applied in the logic governing the rule-based decision matrix of the fraud detection pipeline and the semantic context injection for the LLM.',
    'AI443 Virtual and Augmented Reality (AI Elective): Directly applied in the implementation of the 3D furniture try-on module, converting 2D images to GLB models via Tripo AI and rendering them in physical space using WebAR.'
]
for mod in crmods: bul(doc, mod)

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 – PROBLEM DEFINITION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, 'Problem Definition')
para(doc, 'This chapter defines the precise problem that EzSell addresses, articulates the proposed software solution, lists measurable business objectives, and defines the scope of the system.')

h2(doc, 'Problem Statement')
para(doc, 'The Pakistani second-hand goods market is one of the fastest-growing informal economies, yet it remains riddled with systemic inefficiencies. Existing platforms operate as unstructured boards with no AI integration, automated pricing guidance, or verifiable trust mechanisms.')
para(doc, 'Psychologically, the absence of pricing transparency induces high cognitive load and decision fatigue. Sellers routinely set prices based on intuition or emotional attachment, leading to under-pricing (financial loss) or over-pricing (prolonged listing periods and anxiety). Buyers experience a chronic "trust deficit" and "buyer\'s remorse" anxiety, negotiating blindly without reliable benchmarks.')
para(doc, 'Fraudulent listings—duplicate ads, stolen photos, and scam-baiting descriptions—exploit this lack of structural trust, creating an environment of perpetual suspicion. For furniture buyers specifically, the inability to visualize items in their living space leads to spatial anxiety and costly mismatches.')

h2(doc, 'Problem Solution')
para(doc, 'Many buyers and sellers in the Pakistani second-hand market have requested a structured digital ecosystem that eliminates the guesswork, anxiety, and suspicion inherent in traditional classified advertisement platforms. To address the chronic issue of pricing intuition that leads to financial loss or prolonged sales, the EzSell system implements a dual-track artificial intelligence pricing engine. This engine combines a localized Machine Learning ensemble with real-time Groq LLM market scraping to automatically suggest fair market values, thereby saving sellers time and drastically reducing their cognitive load. By providing these objective pricing benchmarks, the system also protects buyers from inflated prices, subsequently lowering their decision fatigue and purchasing anxiety. Furthermore, to combat the rampant issue of fraudulent listings that exploit the current lack of structural trust, the application deploys an automated, multi-layer fraud prevention pipeline. By seamlessly utilizing zero-shot CLIP image classification and dHash perceptual deduplication upon image upload, the system proactively intercepts stolen photos and category mismatches before they ever reach the public feed. The integration of Natural Language Processing to scan for scam-baiting keywords ensures that the marketplace remains a secure environment, ultimately restoring buyer confidence. For the specific pain points of purchasing second-hand furniture, the future ability for buyers to visualize items in their actual living spaces is achieved through an Augmented Reality try-on module. Utilizing Tripo AI for 3D GLB generation, this feature completely removes spatial anxiety, saving buyers from costly mismatches in size and aesthetic style. Finally, by incorporating a SentenceTransformer semantic recommendation engine, the platform eliminates search fatigue, ensuring that users are presented with highly relevant listings tailored to their behavioral patterns. Collectively, these integrated technologies transform a high-friction, low-trust environment into an intelligent, secure, and psychologically reassuring digital marketplace.')

h2(doc, 'Objectives of the Proposed System')
bul(doc, 'BO-1: Deliver AI-assisted price prediction utilizing an ML ensemble and LLM scraping, achieving 85%+ accuracy within ±20% of actual market prices.')
bul(doc, 'BO-2: Reduce fraudulent listing submissions by at least 70% through a five-stage automated pipeline including image classification and deduplication.')
bul(doc, 'BO-3: Enable immersive, browser-native WebAR furniture visualization without requiring the user to download a dedicated mobile application.')
bul(doc, 'BO-4: Streamline the product discovery process by serving personalized, semantic-based feed recommendations in under 800 milliseconds per query.')
bul(doc, 'BO-5: Facilitate secure and instant buyer-to-seller communication through a real-time, in-app messaging architecture.')

h2(doc, 'Scope')
para(doc, 'The scope of the EzSell project encompasses the design, development, and deployment of a web-based, AI-augmented Customer-to-Customer (C2C) marketplace specifically optimized for the Pakistani economy. The boundaries of the platform are strictly defined to facilitate the trading of three distinct product categories: used mobile phones, laptops, and furniture. Within this domain, the system manages the complete digital lifecycle of a product listing, beginning with a streamlined, category-specific seller onboarding process. As part of its core functionality, the system integrates a dual-track AI pricing engine that utilizes a local machine learning ensemble alongside real-time Large Language Model (LLM) web scraping to generate objective market price estimates. To ensure marketplace integrity, the scope includes an automated, multi-layer fraud prevention pipeline that actively scans all user uploads. This pipeline utilizes zero-shot computer vision classification to verify category alignment, perceptual hashing to block duplicate images, and natural language processing to flag scam-baiting terminology in descriptions. For prospective buyers, the platform\'s range extends to a personalized discovery feed powered by semantic embedding recommendations that match user behavior to relevant listings. A major functionality within the system\'s boundary is the Augmented Reality (AR) try-on module, which converts 2D furniture images into 3D models via a third-party API and renders them directly in the browser\'s spatial viewer. To facilitate negotiations, the scope includes a real-time, in-app messaging architecture and an intelligent chatbot assistant configured to provide contextual marketplace support. While the platform comprehensively manages digital interactions and data validation, it explicitly excludes the processing of in-app financial transactions, meaning all monetary exchanges must occur offline between the respective parties. Furthermore, the boundaries of the proposed solution do not include physical delivery logistics, biometric identity verification, or the provision of dedicated native iOS and Android applications, as the frontend is built entirely as a responsive web application.')

h2(doc, 'Modules')
para(doc, 'The proposed system comprises the following primary functional modules:')

h3(doc, 'Module 1: User Authentication & Account Management')
bul(doc, 'FE-1: Register and login securely using Email/Password with OTP verification.')
bul(doc, 'FE-2: Authenticate instantly using Google OAuth 2.0 integration.')
bul(doc, 'FE-3: Manage active stateless sessions securely via JSON Web Tokens (JWT).')

h3(doc, 'Module 2: Product Listings & Ad Management')
bul(doc, 'FE-1: Create new used-goods listings using category-specific dynamic forms for Mobiles, Laptops, and Furniture.')
bul(doc, 'FE-2: Auto-extract and populate metadata specifications using regex-based text analysis on listing titles.')
bul(doc, 'FE-3: Transition listing statuses through a complete lifecycle including Draft, Active, Hidden, and Sold.')

h3(doc, 'Module 3: Price Prediction')
bul(doc, 'FE-1: Generate a baseline market price utilizing a pre-trained ML ensemble of XGBoost, LightGBM, Random Forest, and Gradient Boosting.')
bul(doc, 'FE-2: Cross-validate baseline predictions with live web scraping of OLX Pakistan via the Groq LLM.')
bul(doc, 'FE-3: Filter scraped pricing data using Interquartile Range (IQR) mathematics to discard extreme outliers and ensure accuracy.')

h3(doc, 'Module 4: Recommendations')
bul(doc, 'FE-1: Track user behavioral metrics including views, clicks, and favorites across different listing categories.')
bul(doc, 'FE-2: Generate a personalized discovery feed by utilizing SentenceTransformer embeddings ranked via cosine similarity algorithms.')

h3(doc, 'Module 5: Analytical Dashboard')
bul(doc, 'FE-1: Provide administrators with Recharts-based visual analytics tracking platform engagement and total active users.')
bul(doc, 'FE-2: Manage a centralized moderation queue allowing admins to manually review, approve, or reject flagged user listings.')

h3(doc, 'Module 6: AR Customization Try-On')
bul(doc, 'FE-1: Convert standard 2D user-uploaded furniture images into photorealistic 3D GLB models using the Tripo AI API.')
bul(doc, 'FE-2: Render 3D furniture models directly within the user\'s physical space utilizing Google model-viewer for spatial WebAR.')

h3(doc, 'Module 7: Fraud Prevention')
bul(doc, 'FE-1: Execute zero-shot CLIP classification on uploaded images to guarantee they match the declared listing category.')
bul(doc, 'FE-2: Identify and block duplicate or stolen photos across the platform utilizing dHash perceptual image hashing.')
bul(doc, 'FE-3: Scan listing descriptions using NLP algorithms to actively flag scam-baiting terminology before a listing is made public.')

h3(doc, 'Module 8: Support & Notifications')
bul(doc, 'FE-1: Manage real-time system-wide alerts and visual unread count badges for all user activities.')
bul(doc, 'FE-2: Provide a dedicated support ticket dashboard allowing users to report issues and administrators to track resolution status.')

h3(doc, 'Module 9: Messaging & Inbox')
bul(doc, 'FE-1: Facilitate secure, real-time WebSocket-based messaging between prospective buyers and sellers.')
bul(doc, 'FE-2: Organize conversations within an inbox view grouped by specific listing to streamline negotiations.')

h3(doc, 'Module 10: Favorites & User Dashboard')
bul(doc, 'FE-1: Allow users to bookmark and save specific active listings into a dedicated favorites list for later viewing.')
bul(doc, 'FE-2: Provide a seller dashboard that tracks live status counts of the user\'s own active, pending, hidden, and sold advertisements.')

h3(doc, 'Module 11: AI Chatbot Assistant')
bul(doc, 'FE-1: Provide an intelligent customer support and market guide powered by the Groq llama-3.1-8b-instant LLM.')
bul(doc, 'FE-2: Stream real-time conversational responses to the client using Server-Sent Events (SSE).')
bul(doc, 'FE-3: Inject live OLX CSV market context directly into the prompt to ensure the chatbot provides accurate platform queries.')


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 – REQUIREMENT ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, 'Requirement Analysis')

h2(doc, 'User classes and characteristics')
para(doc, 'Table 3: User Classes and Characteristics', style='Caption')
simple_table(doc, 
    ['User class', 'Description'], 
    [
        ['Buyer', 'A registered individual looking to purchase second-hand mobiles, laptops, or furniture. They rely heavily on the system\'s AR visualization and semantic recommendations. Buyers interact with the system via the web platform to browse, favorite, and securely message sellers in real time without exposing personal phone numbers.'],
        ['Seller', 'A registered individual looking to liquidate used electronics or furniture. They utilize the AI Price Predictor to set competitive prices and reduce the cognitive load of market research. Sellers interact with the system to create category-specific listings, manage ad statuses (Active, Hidden, Sold), and respond to buyer inquiries.'],
        ['Administrator', 'A system moderator responsible for maintaining platform integrity. Administrators do not buy or sell; instead, they utilize the analytical dashboard to review fraud-flagged listings (via CLIP and dHash algorithms), manage user accounts, resolve support tickets, and track overall marketplace engagement metrics.']
    ]
)

h2(doc, 'Requirement Identifying Technique')
para(doc, 'Use case modeling and event-response tables were selected to map user actions to systemic AI responses. Storyboarding was utilized for the UI/UX design of the AR Viewer and Chatbot.')

h2(doc, 'Functional Requirements')
def fr_table(doc, id_str, title, req_desc, source, rationale, biz_rule, deps, priority):
    rows_data = [
        ('Identifier', id_str),
        ('Title', title),
        ('Requirement', req_desc),
        ('Source', source),
        ('Rationale', rationale),
        ('Business Rule', biz_rule),
        ('Dependencies', deps),
        ('Priority', priority)
    ]
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.style = 'Table Grid'
    for r, (header, val) in enumerate(rows_data):
        t.rows[r].cells[0].text = header
        t.rows[r].cells[1].text = val
        cell_shade(t.rows[r].cells[0], '1F3864')
        for run in t.rows[r].cells[0].paragraphs[0].runs: run.bold=True; run.font.color.rgb=RGBColor(255,255,255)
    doc.add_paragraph()

all_frs = [
    ("FR-1", "Profile Management", "The system shall allow users to register securely via Email or Google OAuth, authenticate via JWT, and manage their profiles.", "User Requirement", "To securely onboard users and manage identity.", "Passwords must be hashed via bcrypt.", "None", "High"),
    ("FR-2", "Product Listings & Management", "The system shall enable sellers to create, edit, hide, and delete category-specific listings with multi-image uploads and AI title validation.", "Core Marketplace", "To populate the platform with structured inventory.", "Must select Mobile, Laptop, or Furniture.", "FR-1", "High"),
    ("FR-3", "Price Prediction", "The system shall calculate a fair market price utilizing a blended ML ensemble (XGBoost, LightGBM, RF, GBC) and live Groq LLM web scraping (Accuracy: 85%+).", "AI Engine", "To reduce seller cognitive load and ensure fair pricing.", "Predictions must be cross-validated against live scraped data.", "FR-2", "High"),
    ("FR-4", "Recommendations", "The system shall generate a personalized discovery feed by encoding listings with SentenceTransformer vectors and ranking via cosine similarity.", "User Experience", "To reduce search fatigue and match user intent.", "Cosine similarity > 0.5 threshold.", "FR-2", "Medium"),
    ("FR-5", "Analytical Dashboard", "The system shall provide administrators with visual Recharts tracking engagement, listings, users, and support tickets for platform moderation.", "Admin Team", "To oversee and moderate marketplace health.", "Requires Admin privileges.", "FR-1", "High"),
    ("FR-6", "AR Customization Try-On", "The system shall invoke the Tripo AI API to generate photorealistic 3D GLB models from images and render them in physical space via WebAR.", "AR Engine", "To eliminate spatial anxiety for furniture buyers.", "Restricted strictly to the Furniture category.", "FR-2", "High"),
    ("FR-7", "Fraud Prevention", "The system shall utilize CLIP zero-shot classification, dHash deduplication, and NLP to proactively block duplicate, stolen, or scam listings.", "Security Team", "To protect marketplace integrity and prevent scams.", "Hashes matched against global database.", "FR-2", "High"),
    ("FR-8", "Support & Notifications", "The system shall facilitate a real-time ticketing dashboard and WebSocket-based notification alerts to keep users informed of status changes.", "Support Team", "To manage customer service workflows.", "Real-time delivery required.", "FR-1", "Medium"),
    ("FR-9", "Messaging & Inbox", "The system shall provide secure, real-time WebSocket messaging grouped by listing, allowing buyers and sellers to negotiate securely.", "Core Marketplace", "To facilitate instant communication without exposing emails.", "Requires active session.", "FR-1", "High"),
    ("FR-10", "Favorites & User Dashboard", "The system shall allow users to bookmark listings and provide sellers with a dashboard to track their active, pending, hidden, and sold items.", "User Experience", "To manage personal inventory and interests.", "None", "FR-2", "Medium"),
    ("FR-11", "AI Chatbot Assistant", "The system shall provide a conversational interface powered by llama-3.1-8b with live context injection to assist users via SSE streaming.", "AI Engine", "To provide automated intelligent support.", "Market data context updated daily.", "None", "Medium"),
]

for idx, fr_data in enumerate(all_frs):
    t_id = fr_data[0]
    t_num = idx + 4 # Start tables from Table 4
    para(doc, f'Table {t_num} Description of {t_id}', style='Caption')
    fr_table(doc, *fr_data)

h2(doc, 'Non-Functional Requirements')
para(doc, 'This section specifies nonfunctional requirements defining system quality attributes. These requirements are specific, quantitative, and verifiable.')

h3(doc, 'Reliability')
bul(doc, 'REL-1: The ML Pricing Engine shall maintain a Mean Time Between Failures (MTBF) of 720 hours. A failure is defined as an inability to generate a price estimate within 10 seconds. In the event of external API rate-limiting, the system shall implement automatic error detection and seamlessly fall back to local cached CSV datasets to prevent service disruption.')

h3(doc, 'Usability')
bul(doc, 'USE-1: The platform shall allow a user to navigate from the homepage to a 3D AR furniture view with a maximum of three interactions (clicks/taps).')
bul(doc, 'USE-2: The listing creation form shall implement real-time inline validation to prevent submission errors, guiding the user to intuitively correct missing category-specific attributes.')

h3(doc, 'Performance')
bul(doc, 'PER-1: 95% of personalized recommendation feeds generated by the SentenceTransformer model shall load completely within 800 milliseconds from the time of the user request over a standard 4G connection.')
bul(doc, 'PER-2: The 3D GLB model generation process must initiate browser rendering within 45 seconds of the user invoking the AR Try-On feature.')

h3(doc, 'Security')
bul(doc, 'SEC-1: The system shall protect all user personal identifiable information (PII) and chat logs against unauthorized access, requiring advanced administrative credentials and secure cryptographic verification to access the central database.')
bul(doc, 'SEC-2: The application shall prevent automated bot scraping of seller contact information, ensuring phone numbers are completely obscured from unauthenticated entities.')

h2(doc, 'External Interface Requirements')
para(doc, 'This section provides information ensuring the system communicates properly with users and external software/hardware elements.')

h3(doc, 'User Interfaces Requirements')
bul(doc, 'UI-1: The web application shall adhere strictly to the Shadcn UI product family style guidelines for all form controls, alerts, and typography.')
bul(doc, 'UI-2: The interface shall implement a fully responsive grid system to accommodate standard mobile (375px width) and desktop (1440px width) screen resolution constraints.')
bul(doc, 'UI-3: A fixed navigation bar shall appear on every screen, providing instant access to the user inbox, favorites, and a floating action button for the AI Chatbot.')

h3(doc, 'Software Interfaces')
para(doc, 'SI-1: AI Inference Services')
bul(doc, 'SI-1.1: The backend FastAPI service shall transmit uploaded listing images to the OpenAI CLIP model through a programmatic REST interface for category validation prior to database insertion.')
para(doc, 'SI-2: Tripo AI 3D Generation')
bul(doc, 'SI-2.1: The AR Engine shall send 2D image binaries to the Tripo AI API over HTTP and receive processed GLB 3D model URLs for browser rendering.')

h3(doc, 'Hardware Interfaces')
bul(doc, 'HI-1: The frontend Single Page Application (SPA) shall request permissions and interface directly with the user\'s mobile device camera hardware via WebXR and Google model-viewer to project AR models into the physical environment.')

h3(doc, 'Communications Interfaces')
bul(doc, 'CI-1: The system shall send an automated transactional email containing a 6-digit OTP to the user via SMTP protocols to confirm account activation.')
bul(doc, 'CI-2: The real-time messaging module shall utilize secure WebSocket (WSS) network protocols to transmit chat packets instantly between buyers and sellers.')

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 – DESIGN AND ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, 'Design and Architecture')

h2(doc, 'Architectural Design')
para(doc, 'EzSell follows a Multi-Tiered Client-Server Architecture. Frontend (React SPA) connects to Application Logic (FastAPI), which interfaces with Data/AI (PostgreSQL, ML models, External APIs).')
placeholder(doc, 'Box and Line Architecture Diagram')

h2(doc, 'Design Models')
h3(doc, 'Activity Diagram: Listing Creation Flow')
para(doc, 'Maps workflow of a user creating a new listing, illustrating concurrent execution of CLIP validation and AI Price Prediction.')
placeholder(doc, 'Activity Diagram: Listing Creation')

h3(doc, 'Class Diagram: ORM Schema')
para(doc, 'Represents SQLAlchemy ORM. Highlights polymorphic relationships between base Listing and specialized Mobile/Laptop/Furniture entities.')
placeholder(doc, 'Class Diagram: ORM')

h2(doc, 'Data Design')
h3(doc, 'Data Dictionary')
para(doc, 'The EzSell platform utilizes an Object-Oriented (OO) data model mapped via SQLAlchemy ORM. The following table alphabetically lists the primary core entities, their attributes, methods, and descriptive purpose.')

para(doc, 'Table 15: Object-Oriented Data Dictionary', style='Caption')
simple_table(doc,
    ['Object / Entity', 'Attributes (Name : Type)', 'Methods (Name(Params))', 'Description'],
    [
        ['Conversation', 'id: UUID\nlisting_id: UUID\nbuyer_id: UUID', 'get_messages(limit: Int)\nmark_read(user_id: UUID)', 'Groups messages between a buyer and seller for a specific listing.'],
        ['Favorite', 'id: Int\nuser_id: UUID\nlisting_id: UUID\ncreated_at: DateTime', 'None', 'Association object tracking users\' bookmarked listings.'],
        ['Listing', 'id: UUID\ntitle: String\ndescription: Text\nprice: Float\nstatus: Enum\npredicted_price: Float\ncategory: Enum\nimages: JSON\nspecs: JSON', 'validate_ai_price()\nupdate_status(new_status: Enum)\ncheck_fraud_hash()', 'Core inventory item (Mobile, Laptop, Furniture) representing an ad.'],
        ['Message', 'id: Int\nconv_id: UUID\nsender_id: UUID\ncontent: Text\nread: Boolean\ntimestamp: DateTime', 'None', 'Individual plaintext chat messages transmitted via WebSockets.'],
        ['User', 'id: UUID\nemail: String\npassword_hash: String\nphone: String\nrole: Enum\navatar_url: String\nis_active: Boolean', 'verify_otp(code: String)\ngenerate_jwt() -> String\nupdate_profile(data: Dict)', 'System actor representing a buyer, seller, or administrator.']
    ]
)

h2(doc, 'Human Interface Design')
h3(doc, 'Screen Images')
placeholder(doc, 'Screenshot: Home Feed')
placeholder(doc, 'Screenshot: AR Viewer Overlay')

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 – IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, 'Implementation')

h2(doc, 'Algorithm')
para(doc, 'This section outlines the detailed pseudocode logic driving the major intelligent modules of the system. The algorithms are represented with standard programmatic semantics (e.g., assignment, loops, conditionals) to accurately reflect their operation.')

h3(doc, 'Module 1: User Authentication & Account Management')
para(doc, 'Table 21: Algorithm 1 - User Authentication (JWT Generation)', style='Caption')
algo_table(doc, 'Algorithm 1: JWT_Generation', 'User Credentials (email, password)', 'JWT Token (string) or Error', [
  "1: user_record ← DB_Find_User(email)",
  "2: if user_record == null then",
  "3:     return Error('Invalid credentials')",
  "4: end if",
  "5: is_valid ← bcrypt.verify(password, user_record.password_hash)",
  "6: if is_valid == false then",
  "7:     return Error('Invalid credentials')",
  "8: end if",
  "9: payload ← { 'sub': user_record.id, 'exp': current_time() + 24_hours }",
  "10: token ← HS256_Sign(payload, SECRET_KEY)",
  "11: return token"
])

h3(doc, 'Module 2: Product Listings & Ad Management')
para(doc, 'Table 22: Algorithm 2 - Listing Category Inference', style='Caption')
algo_table(doc, 'Algorithm 2: Metadata_Inference', 'Title (string), Description (string)', 'Extracted Metadata (Dictionary)', [
  "1: Metadata ← {}",
  "2: title_tokens ← split(lower(Title), ' ')",
  "3: foreach token in title_tokens do:",
  "4:     if token matches Regex('[0-9]+GB') then",
  "5:         Metadata['RAM'] ← token",
  "6:     end if",
  "7:     if token matches Regex('(iphone|samsung|dell|hp)') then",
  "8:         Metadata['Brand'] ← extract_brand(token)",
  "9:     end if",
  "10: end foreach",
  "11: return Metadata"
])

h3(doc, 'Module 3: AI Price Prediction Pipeline')
para(doc, 'Table 23: Algorithm 3 - AI Price Prediction', style='Caption')
algo_table(doc, 'Algorithm 3: AI_Price_Prediction_Pipeline', 'Listing object L with attributes (title, specs, category)', 'Final Predicted Price (Final_Price), Confidence Score (Conf)', [
  "1: weights ← {XGB: 0.35, LGBM: 0.35, RF: 0.15, GBC: 0.15}",
  "2: ML_Base ← (weights.XGB * Predict_XGB(L)) + (weights.LGBM * Predict_LGBM(L)) + (weights.RF * Predict_RF(L)) + (weights.GBC * Predict_GBC(L))",
  "3: query_string ← concatenate('site:olx.com.pk ', L.category, ' ', L.title)",
  "4: Raw_URLs ← DuckDuckGo_Search(query_string, max_results=10)",
  "5: Scraped_Prices ← []",
  "6: foreach url in Raw_URLs do:",
  "7:     price ← Groq_LLM_Extract_Price(Scrape_HTML(url))",
  "8:     if price is valid then",
  "9:         Scraped_Prices.append(price)",
  "10:     end if",
  "11: end foreach",
  "12: if len(Scraped_Prices) < 3 then",
  "13:     Final_Price ← ML_Base",
  "14:     Conf ← 0.60",
  "15:     return Final_Price, Conf",
  "16: end if",
  "17: Q1 ← Calculate_Percentile(Scraped_Prices, 25)",
  "18: Q3 ← Calculate_Percentile(Scraped_Prices, 75)",
  "19: IQR ← Q3 - Q1",
  "20: Lower_Bound ← Q1 - (1.5 * IQR)",
  "21: Upper_Bound ← Q3 + (1.5 * IQR)",
  "22: Filtered_Prices ← []",
  "23: foreach p in Scraped_Prices do:",
  "24:     if (p ≥ Lower_Bound and p ≤ Upper_Bound) then",
  "25:         Filtered_Prices.append(p)",
  "26:     end if",
  "27: end foreach",
  "28: LLM_Market ← Average(Filtered_Prices)",
  "29: Final_Price ← (ML_Base * 0.85) + (LLM_Market * 0.15)",
  "30: Conf ← Calculate_Confidence(Variance(Filtered_Prices), ML_Base, LLM_Market)",
  "31: return Final_Price, Conf"
])

h3(doc, 'Module 4: Semantic Recommendations')
para(doc, 'Table 24: Algorithm 4 - Semantic Listing Ranking', style='Caption')
algo_table(doc, 'Algorithm 4: Semantic_Ranking', 'User interest vector U, Set of all active listings L', 'Ranked list of recommended listings R', [
  "1: R ← []",
  "2: threshold ← 0.5",
  "3: foreach item in L do:",
  "4:     if item.category == U.primary_category then",
  "5:         item_vector ← SentenceTransformer_Encode(item.title + ' ' + item.description)",
  "6:         similarity_score ← 1 - spatial.distance.cosine(U.vector, item_vector)",
  "7:         if similarity_score ≥ threshold then",
  "8:             R.append({listing: item, score: similarity_score})",
  "9:         end if",
  "10:     end if",
  "11: end foreach",
  "12: R ← Sort_Descending_By_Score(R)",
  "13: return R.limit(20)"
])

h3(doc, 'Module 5: Analytical Dashboard')
para(doc, 'Table 25: Algorithm 5 - Metrics Aggregation', style='Caption')
algo_table(doc, 'Algorithm 5: Metrics_Aggregation', 'Date Range (start_date, end_date)', 'Aggregated Statistics Object', [
  "1: Stats ← {}",
  "2: Stats.Total_Users ← DB_Count('Users', where='created_at BETWEEN start_date AND end_date')",
  "3: Stats.Active_Listings ← DB_Count('Listings', where='status == Active')",
  "4: Stats.Pending_Tickets ← DB_Count('Tickets', where='status == Open')",
  "5: Stats.Daily_Signups ← DB_Group_By_Day('Users', start_date, end_date)",
  "6: return Stats"
])

h3(doc, 'Module 6: AR Customization Try-On')
para(doc, 'Table 26: Algorithm 6 - 3D GLB Model Generation', style='Caption')
algo_table(doc, 'Algorithm 6: AR_3D_Model_Generation', 'List of 2D images I, Category C', 'GLB Model URL or Error', [
  "1: if C != 'Furniture' then",
  "2:     return Error('AR Try-On restricted to Furniture category.')",
  "3: end if",
  "4: Primary_Image ← Select_Highest_Resolution(I)",
  "5: Processed_Image ← Remove_Background_And_Center(Primary_Image)",
  "6: Task_ID ← Tripo_API_Submit(Processed_Image)",
  "7: status ← 'Processing'",
  "8: timeout ← 60 seconds",
  "9: while (status == 'Processing' and timeout > 0) do:",
  "10:    Wait(2 seconds)",
  "11:    timeout ← timeout - 2",
  "12:    status, Model_URL ← Tripo_API_Check_Status(Task_ID)",
  "13:    if status == 'Success' then",
  "14:        return Model_URL",
  "15:    else if status == 'Failed' then",
  "16:        return Error('3D Generation Failed.')",
  "17:    end if",
  "18: end while",
  "19: return Error('Timeout exceeded.')"
])

h3(doc, 'Module 7: Fraud Prevention Pipeline')
para(doc, 'Table 27: Algorithm 7 - Fraud Detection and Blocking', style='Caption')
algo_table(doc, 'Algorithm 7: Fraud_Prevention_Pipeline', 'Listing L, Array of Images Imgs', 'Boolean status (True = Approved, False = Rejected/Flagged)', [
  "1: Listing_Hash ← SHA256(L.title + L.price + L.category)",
  "2: if (Database_Contains(Listing_Hash) == true) then",
  "3:     Mark_As_Rejected(L, 'Duplicate Listing')",
  "4:     return False",
  "5: end if",
  "6: foreach img in Imgs do:",
  "7:     pred_class ← CLIP_Zero_Shot_Predict(img, ['mobile', 'laptop', 'furniture', 'random'])",
  "8:     if (pred_class != L.category) then",
  "9:         Mark_As_Rejected(L, 'Image Category Mismatch')",
  "10:        return False",
  "11:    end if",
  "12:    img_hash ← calculate_dHash(img)",
  "13:    matches ← Find_Similar_Hashes(img_hash, threshold=5)",
  "14:    if (len(matches) > 0) then",
  "15:        Mark_As_Rejected(L, 'Stolen/Reused Image')",
  "16:        return False",
  "17:    end if",
  "18: end foreach",
  "19: Scam_Keywords ← ['send money', 'western union', 'advance payment']",
  "20: foreach word in Scam_Keywords do:",
  "21:    if Regex_Match(L.description, word) then",
  "22:        Mark_As_Flagged(L, 'Manual Review Required')",
  "23:        return False",
  "24:    end if",
  "25: end foreach",
  "26: Mark_As_Approved(L)",
  "27: return True"
])

h3(doc, 'Module 8: Support & Notifications')
para(doc, 'Table 28: Algorithm 8 - Support Ticket Creation', style='Caption')
algo_table(doc, 'Algorithm 8: Ticket_Creation', 'User U, Subject (string), Description (string)', 'Ticket ID', [
  "1: if len(Description) < 10 then",
  "2:     return Error('Description too short')",
  "3: end if",
  "4: Ticket ← Create_DB_Record('Tickets', user_id=U.id, subject=Subject, desc=Description, status='Open')",
  "5: trigger_event('ticket_created', Ticket)",
  "6: return Ticket.id"
])

h3(doc, 'Module 9: Messaging & Inbox')
para(doc, 'Table 29: Algorithm 9 - WebSocket Message Routing', style='Caption')
algo_table(doc, 'Algorithm 9: WebSocket_Routing', 'Sender_ID, Receiver_ID, Message_Content, Listing_ID', 'Boolean delivery status', [
  "1: Conv_ID ← get_or_create_conversation(Sender_ID, Receiver_ID, Listing_ID)",
  "2: Msg ← store_message_in_db(Conv_ID, Sender_ID, Message_Content)",
  "3: if is_user_connected(Receiver_ID) then",
  "4:     socket ← get_active_websocket(Receiver_ID)",
  "5:     send_payload(socket, Msg)",
  "6:     Msg.delivered ← true",
  "7: else",
  "8:     increment_unread_badge(Receiver_ID)",
  "9: end if",
  "10: return Msg.delivered"
])

h3(doc, 'Module 10: Favorites & User Dashboard')
para(doc, 'Table 30: Algorithm 10 - Favorites Toggle Logic', style='Caption')
algo_table(doc, 'Algorithm 10: Favorites_Toggle', 'User_ID, Listing_ID', 'New State ("Added" or "Removed")', [
  "1: fav_record ← DB_Query('Favorites', user_id=User_ID, listing_id=Listing_ID)",
  "2: if fav_record exists then",
  "3:     DB_Delete(fav_record)",
  "4:     return 'Removed'",
  "5: else",
  "6:     DB_Insert('Favorites', user_id=User_ID, listing_id=Listing_ID)",
  "7:     return 'Added'",
  "8: end if"
])

h3(doc, 'Module 11: AI Chatbot Assistant')
para(doc, 'Table 31: Algorithm 11 - Context Injection', style='Caption')
algo_table(doc, 'Algorithm 11: AI_Context_Injection', 'User Query Q', 'Streamed Response Generator', [
  "1: market_data ← Load_CSV('olx_market_summary.csv')",
  "2: system_prompt ← 'You are EzSell Assistant. Use this data: ' + format_to_string(market_data)",
  "3: full_prompt ← system_prompt + '\\nUser: ' + Q",
  "4: stream ← Groq_API_Call('llama-3.1-8b-instant', full_prompt, stream=True)",
  "5: foreach token in stream do:",
  "6:     yield token",
  "7: end foreach"
])

h2(doc, 'External APIs/SDKs')
para(doc, 'Describe the third-party APIs/SDKs used in the project implementation in the following table.')
para(doc, 'Table 32 Details of APIs used in the project', style='Caption')
simple_table(doc,
    ['Name of API and version', 'Description of API', 'Purpose of usage', 'List down the API endpoint/function/class in which it is used'],
    [
        ['Groq API (v1)', 'Ultra-fast LLM inference API', 'llama-3.3-70b for price validation; llama-3.1-8b for chatbot streaming.', 'https://api.groq.com/openai/v1/chat/completions'],
        ['Tripo AI (Version V2)', 'Image-to-3D generation service', 'Generates GLB 3D models from furniture photos for AR try-on.', 'https://api.tripo3d.ai/v2/openapi/task'],
        ['Cloudinary API (v1_1)', 'Cloud image and video management solution', 'Uploading Product Images on Cloudinary server', 'https://api.cloudinary.com/v1_1/demo/image/upload'],
        ['DuckDuckGo DDGS (Latest)', 'Python search scraping library', 'Scrapes live OLX Pakistan links to retrieve recent market prices.', 'duckduckgo_search.DDGS.text'],
        ['Google OAuth 2.0 (v2)', 'Authentication and authorization framework', 'Enables users to securely log in using their Google accounts.', 'https://oauth2.googleapis.com/token']
    ]
)

h2(doc, 'User Interface')
para(doc, 'The React frontend uses Shadcn UI and Tailwind CSS. The Home Feed implements infinite scrolling. The AR Viewer uses Google model-viewer for native spatial visualization.')

h2(doc, 'Deployment')
para(doc, 'During development, we ran and tested the whole system on our local machines. The project had two main parts: the backend API and the frontend web application. Both were connected and tested together in a local setup.')

h3(doc, 'Backend Environment')
bul(doc, 'Framework: FastAPI')
bul(doc, 'Runtime: Python 3.13')
bul(doc, 'Server Used: Uvicorn (development server)')
bul(doc, 'Hosting Mode: Localhost (http://127.0.0.1:8000/)')
bul(doc, 'Main Dependencies: fastapi, uvicorn, sqlalchemy, pydantic, httpx, python-jose, passlib[bcrypt], google-genai, python-multipart, python-dotenv, pandas, numpy, matplotlib, prophet, transformers, torch, groq, bcrypt')

h3(doc, 'Frontend Environment')
bul(doc, 'Framework: React.js')
bul(doc, 'Runtime Environment: Node.js (LTS version)')
bul(doc, 'Local Hosting URL: http://localhost:8080/')
bul(doc, 'Package Manager: npm')
bul(doc, 'Primary Dependencies: React Router, Axios, Redux / Context API, Tailwind CSS, Utility Libraries')

h3(doc, 'Testing Setup')
bul(doc, 'Simultaneous Local Execution: Both the FastAPI backend and the React.js frontend were run together on the local development system.')
bul(doc, 'API Communication: The frontend interacted with the backend using REST API calls over local URLs.')
bul(doc, 'Modules Tested in Local Environment: User Authentication, Product Listings, AI Price Prediction, Dashboards, 3D AR Module, Database CRUD Operations.')

para(doc, 'For production, the EzSell platform is deployed on Amazon Web Services (AWS) to guarantee high availability, low-latency access, and robust computational performance.')

h3(doc, 'Cloud Infrastructure Details')
simple_table(doc,
    ['Specification', 'Detail'],
    [
        ['Hosting Service', 'Amazon Web Services (AWS) Elastic Compute Cloud (EC2)'],
        ['Instance ID', 'i-037ea3edb54596c83'],
        ['Instance Name', 'ezsell backend'],
        ['Region & Availability Zone', 'eu-north-1 (Europe - Stockholm) / eu-north-1c'],
        ['Status', 'Running']
    ]
)

h3(doc, 'Compute and Storage Specifications')
simple_table(doc,
    ['Component', 'Specification Details'],
    [
        ['Instance Type', 'c7i-flex.large (x86_64 Architecture)'],
        ['CPU', '1 core with 2 threads per core'],
        ['Operating System', 'Ubuntu 26.04 LTS (ubuntu-resolute-26.04-amd64-server-20260421)'],
        ['Storage Volume', '16 GiB gp3 SSD (Root Volume)'],
        ['Storage Performance', '3,000 IOPS, 125 MB/s Throughput'],
        ['Encryption', 'Not encrypted, Delete on Termination Enabled']
    ]
)

h3(doc, 'Network and Security Configuration')
simple_table(doc,
    ['Parameter', 'Configuration Details'],
    [
        ['VPC and IPs', 'Default VPC (vpc-092755086093495e8) | Private IP: 172.31.0.165 | Public IP: 13.51.45.172'],
        ['Public DNS', 'ec2-13-51-45-172.eu-north-1.compute.amazonaws.com'],
        ['Security Group', 'launch-wizard-1'],
        ['Inbound Rules', 'Port 22 (SSH), Port 80 (HTTP), Port 443 (HTTPS) - Open to all IPs (0.0.0.0/0)'],
        ['Outbound Rules', 'All outbound traffic allowed']
    ]
)

h3(doc, 'Access Information and Software Stack')
bul(doc, 'SSH Access Command: ssh -i ezsell.pem ubuntu@13.51.45.172')
bul(doc, 'Web Server Access: http://13.51.45.172 or https://13.51.45.172')
bul(doc, 'Web Server Software: Uvicorn ASGI Server managed by Gunicorn process workers.')


# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 8 – TESTING AND EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, 'Testing and Evaluation')

h2(doc, 'API & Module Endpoint Testing')
para(doc, 'The following tables detail the endpoint-level test cases and their execution outcomes across the core modules of the application.')

h3(doc, 'M1 — Profile Management Tests')
para(doc, 'Table 33: M1 Profile Management Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M1-1', 'Register new user with email & phone', 'POST /auth/register', '201 Created, user in DB', 'Pass'],
    ['T-M1-2', 'Login with correct credentials', 'POST /auth/login', '200 OK, JWT returned', 'Pass'],
    ['T-M1-3', 'Login with wrong password', 'POST /auth/login', '401 Unauthorized', 'Pass'],
    ['T-M1-4', 'Google OAuth sign-in flow', 'Browser OAuth redirect', 'Token set, user created/fetched', 'Pass'],
    ['T-M1-5', 'Access secured endpoint without token', 'GET /users/me', '401 Unauthorized', 'Pass'],
    ['T-M1-6', 'Update profile fields', 'PATCH /users/me', '200 OK, fields updated', 'Pass'],
    ['T-M1-7', 'Upload avatar image', 'POST /users/me/avatar', '200 OK, avatar_url set', 'Pass'],
    ['T-M1-8', 'Admin flag verification', 'GET /users/me as admin', 'is_admin: true in response', 'Pass']
])

h3(doc, 'M2 — Product Listings Tests')
para(doc, 'Table 34: M2 Product Listings Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M2-1', 'Create listing with 3 images', 'POST /listings/', '201 Created', 'Pass'],
    ['T-M2-2', 'Fetch listing by ID', 'GET /listings/{id}', '200 OK, full listing data', 'Pass'],
    ['T-M2-3', 'Edit listing title', 'PATCH /listings/{id}', '200 OK, title updated', 'Pass'],
    ['T-M2-4', 'Delete listing as owner', 'DELETE /listings/{id}', '200 OK, removed from DB', 'Pass'],
    ['T-M2-5', 'Hide listing (toggle)', 'PATCH /listings/{id}/toggle-visibility', 'is_active: false', 'Pass'],
    ['T-M2-6', 'Mark listing as sold', 'PATCH /listings/{id}/mark-sold', 'is_sold: true', 'Pass'],
    ['T-M2-7', 'Submit listing for approval', 'POST /listings/', 'approval_status: pending', 'Pass'],
    ['T-M2-8', 'Admin approve listing', 'PATCH /approvals/{id}/approve', 'approval_status: approved', 'Pass'],
    ['T-M2-9', 'Search by category', 'GET /listings/?category=mobile', 'Filtered results returned', 'Pass'],
    ['T-M2-10', 'AI title validation', 'POST /listings/validate-title', 'Valid/invalid flag returned', 'Pass']
])

h3(doc, 'M3 — Price Prediction Tests')
para(doc, 'Table 35: M3 Price Prediction Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M3-1', 'Predict price for mobile phone (ML)', 'POST /predictions/mobile', 'Price range + confidence score', 'Pass'],
    ['T-M3-2', 'Predict price for laptop (ML)', 'POST /predictions/laptop', 'Price range + confidence score', 'Pass'],
    ['T-M3-3', 'Predict price for furniture (ML)', 'POST /predictions/furniture', 'Price range + confidence score', 'Pass'],
    ['T-M3-4', 'Groq LLM price estimation', 'POST /predictions/advanced', 'LLM-sourced price + OLX snippets', 'Pass'],
    ['T-M3-5', 'OLX DuckDuckGo scraping', 'Internal service call', 'Prices extracted and IQR-filtered', 'Pass'],
    ['T-M3-6', 'Prediction on create form', 'Frontend UX', 'Suggested price shown to seller', 'Pass'],
    ['T-M3-7', 'Confidence score < 50% flagged', 'Frontend UX', 'Low confidence warning shown', 'Pass']
])

h3(doc, 'M4 — Recommendations Tests')
para(doc, 'Table 36: M4 Recommendations Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M4-1', 'Log view activity', 'POST /recommendations/activity', '200 OK, activity recorded', 'Pass'],
    ['T-M4-2', 'Get personalized feed', 'GET /recommendations/', 'Returns interest-matched listings', 'Pass'],
    ['T-M4-3', 'Get trending for anonymous', 'GET /recommendations/trending', 'Returns top-viewed listings', 'Pass'],
    ['T-M4-4', 'Filter by price range', 'GET /listings/?min_price=X&max_price=Y', 'Listings within range returned', 'Pass'],
    ['T-M4-5', 'Semantic similarity search', 'GET /listings/?q=leather+sofa', 'NLP-ranked results returned', 'Pass']
])

h3(doc, 'M5 — Analytical Dashboard Tests')
para(doc, 'Table 37: M5 Analytical Dashboard Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M5-1', 'Admin fetch dashboard stats', 'GET /analytics/overview', 'User/listing/revenue totals', 'Pass'],
    ['T-M5-2', 'Admin view all users', 'GET /users/ (admin)', 'Full user list with details', 'Pass'],
    ['T-M5-3', 'Admin view all listings', 'GET /listings/ (admin)', 'All listings regardless of status', 'Pass'],
    ['T-M5-4', 'Admin reject listing with reason', 'PATCH /approvals/{id}/reject', 'Status set, rejection reason saved', 'Pass'],
    ['T-M5-5', 'Admin view support tickets', 'GET /support/admin/tickets', 'All tickets with user info', 'Pass'],
    ['T-M5-6', 'Support requests count card', 'Frontend UX', 'Correct total + open count shown', 'Pass']
])

h3(doc, 'M6 — AR / Try-On Tests')
para(doc, 'Table 38: M6 AR / Try-On Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M6-1', 'Generate procedural 3D model', 'POST /products/{id}/assets/generate', 'GLB/USDZ URLs returned', 'Pass'],
    ['T-M6-2', 'Launch WebAR viewer', 'Frontend UX (/ar/{id})', 'model-viewer loads GLB', 'Pass'],
    ['T-M6-3', 'Trigger Tripo AI generation', 'POST /products/{id}/assets/generate-3d', 'Task ID returned, polling starts', 'Pass'],
    ['T-M6-4', 'Poll AI generation status', 'GET /products/{id}/assets/generate-3d/{task_id}', 'status: done, URLs available', 'Pass'],
    ['T-M6-5', 'Admin manual GLB upload', 'POST /products/{id}/assets/upload-glb', 'GLB URL stored in DB', 'Pass'],
    ['T-M6-6', 'iOS USDZ AR QuickLook', 'iOS Safari', 'Native AR viewer launched', 'Pass'],
    ['T-M6-7', 'Fixed scale lock', 'Frontend UX', 'Model does not resize freely', 'Pass']
])

h3(doc, 'M7 — Fraud Prevention Tests')
para(doc, 'Table 39: M7 Fraud Prevention Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M7-1', 'Detect duplicate listing', 'POST duplicate listing', 'listing_hash collision detected', 'Pass'],
    ['T-M7-2', 'Detect duplicate image', 'POST listing with reused image', 'image_hash collision flagged', 'Pass'],
    ['T-M7-3', 'Fraud flags stored', 'DB inspection', 'JSON array in fraud_flags column', 'Pass'],
    ['T-M7-4', 'Flagged listing surfaced to admin', 'Admin Dashboard', 'Flagged items appear in review queue', 'Pass'],
    ['T-M7-5', 'Reject with documented reason', 'PATCH /approvals/{id}/reject + reason', 'rejection_reason field saved in DB', 'Pass']
])

h3(doc, 'M8 — Support & Notifications Tests')
para(doc, 'Table 40: M8 Support & Notifications Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M8-1', 'Submit support ticket', 'POST /support/tickets', '201 Created, ticket in DB', 'Pass'],
    ['T-M8-2', 'Submit bug report', 'POST /support/tickets (type: bug)', '201 Created', 'Pass'],
    ['T-M8-3', 'Admin fetch all tickets', 'GET /support/admin/tickets', 'All tickets with user info returned', 'Pass'],
    ['T-M8-4', 'Admin update ticket to working', 'PATCH /support/admin/tickets/{id}/status', 'Status updated, notification created', 'Pass'],
    ['T-M8-5', 'Admin update ticket to done', 'PATCH /support/admin/tickets/{id}/status', 'Status updated, notification created', 'Pass'],
    ['T-M8-6', 'User fetch notifications', 'GET /notifications', 'Notification list returned', 'Pass'],
    ['T-M8-7', 'Unread count endpoint', 'GET /notifications/unread/count', 'Correct count returned', 'Pass'],
    ['T-M8-8', 'Mark notification as read', 'POST /notifications/{id}/read', 'is_read: true set in DB', 'Pass'],
    ['T-M8-9', 'Mark all notifications as read', 'POST /notifications/read-all', 'All user notifications marked read', 'Pass'],
    ['T-M8-10', 'Bell badge shows unread count', 'Frontend UX', 'Badge visible with correct number', 'Pass'],
    ['T-M8-11', 'Notification popover renders', 'Frontend UX', 'Alert list visible in dropdown', 'Pass'],
    ['T-M8-12', 'Click notification navigates user', 'Frontend UX', 'User routed to /profile', 'Pass']
])

h3(doc, 'M9 — Messaging Tests')
para(doc, 'Table 41: M9 Messaging Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M9-1', 'Send message to seller', 'POST /messages/', '201 Created, stored in DB', 'Pass'],
    ['T-M9-2', 'Fetch conversations list', 'GET /messages/conversations', 'All threads returned', 'Pass'],
    ['T-M9-3', 'Read messages in a conversation', 'GET /messages/{conversation_id}', 'Message history returned', 'Pass'],
    ['T-M9-4', 'Unread message count in nav', 'GET /messages/unread/count', 'Correct count in badge', 'Pass'],
    ['T-M9-5', 'Mark conversation as read', 'POST /messages/{id}/read', 'is_read: true for all messages', 'Pass']
])

h3(doc, 'M10 — Favorites & Dashboard Tests')
para(doc, 'Table 42: M10 Favorites & Dashboard Tests', style='Caption')
simple_table(doc, ['Test ID', 'Test Case', 'Method / Layer', 'Expected Result', 'Outcome'], [
    ['T-M10-1', 'Add listing to favorites', 'POST /favorites/{id}', '200 OK, saved in DB', 'Pass'],
    ['T-M10-2', 'Remove listing from favorites', 'DELETE /favorites/{id}', '200 OK, removed from DB', 'Pass'],
    ['T-M10-3', 'Fetch all user favorites', 'GET /favorites/', 'Full favorites list returned', 'Pass'],
    ['T-M10-4', 'Check if listing is favorited', 'GET /favorites/{id}/check', 'is_favorited: true/false', 'Pass'],
    ['T-M10-5', 'Dashboard shows own listings', 'GET /listings/my', 'Only user\'s listings returned', 'Pass'],
    ['T-M10-6', 'Dashboard status count cards', 'Frontend UX', 'Correct active/hidden/sold counts', 'Pass']
])

h2(doc, 'Unit Testing')
para(doc, 'Unit testing checks individual components to ensure each one behaves as expected.')

h3(doc, 'Unit Testing 1: Validate login functionality with correct and incorrect inputs.')
para(doc, 'Testing Objective: To ensure the login form is working correctly with valid and invalid inputs.')
para(doc, 'Table 43: Unit Testing 1 login Functionalities', style='Caption')
simple_table(doc, ['No.', 'Test case', 'Attribute and value', 'Expected result', 'Result'], [
    ['1', 'Check valid login', 'username: ahmed\nPassword: ahmed123', 'User is logged in and redirected to dashboard.', 'Pass'],
    ['2', 'Invalid email format', 'username: ahme', 'Error, no such sort of username', 'Pass'],
    ['3', 'Invalid password', 'Password: EZsell123', 'Dashboard unaccessible, login page refresh', 'Pass']
])

h3(doc, 'Unit Testing 2: Listing Image Upload')
para(doc, 'Testing Objective: Ensure supported file types are accepted and invalid formats are rejected.')
para(doc, 'Table 44: Unit Testing 2 Listing Image Upload', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Result'], [
    ['1', 'Upload supported image', 'image.jpg', 'File accepted and uploaded to storage', 'Pass'],
    ['2', 'Upload unsupported format', 'file.txt', 'Wont show any other sort of file other than jpg, jpeg, png', 'Pass'],
    ['3', 'Upload large file', 'big_image.webp (10MB)', 'Error', 'Pass']
])

h3(doc, 'Unit Testing 3: Price Prediction')
para(doc, 'Testing Objective: Ensure that the AI price prediction module correctly processes features and returns a valid predicted price with confidence.')
para(doc, 'Table 45: Unit Testing 3 Price Prediction', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Result'], [
    ['1', 'Predict price for mobile', 'title=Samsung s23, condition=used', 'Return predicted price + confidence', 'Pass'],
    ['2', 'Predict price for laptop', 'brand=HP omen, core i3, RAM=8GB', 'Return predicted price + confidence', 'Pass'],
    ['3', 'Missing required feature', 'condition = null', 'No Prediction', 'Pass']
])

h3(doc, 'Unit Testing 4: Listing Creation (Post Ad)')
para(doc, 'Testing Objective: Ensure that new listings are created successfully with all required fields.')
para(doc, 'Table 46: Unit Testing 4 Listing Creation', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Result'], [
    ['1', 'Valid listing', 'title, desc, price, category, image provided', 'Listing saved successfully', 'Pass'],
    ['2', 'Missing title', 'title=null', 'Error', 'Pass'],
    ['3', 'Price not numeric', 'price="ten thousand"', 'Error', 'Pass'],
    ['4', 'Image missing', 'no images uploaded', 'Error', 'Pass']
])

h3(doc, 'Unit Testing 5: User Dashboard Data Fetch')
para(doc, 'Testing Objective: Ensure that dashboard fetches correct analytics based on user_id.')
para(doc, 'Table 47: Unit Testing 5 Dashboard Data', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Result'], [
    ['1', 'Fetch dashboard data', 'user_id valid', 'Returns totals (views, listings, contacts)', 'Pass'],
    ['2', 'No listings', 'user with 0 listings', 'Show zero analytics', 'Pass'],
    ['3', 'Invalid user_id', 'user_id = null', 'Error', 'Pass']
])

h3(doc, 'Unit Testing 6: Admin Analytics Dashboard')
para(doc, 'Testing Objective: Verify correct aggregation of platform-wide metrics.')
para(doc, 'Table 48: Unit Testing 6 Admin Analytics', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Result'], [
    ['1', 'Fetch system metrics', 'none', 'Return totals: users, listings, active listings', 'Pass'],
    ['2', 'No users available', 'empty database', 'Display zero metrics, no errors', 'Pass']
])

h3(doc, 'Unit Testing 7: AR Room Analysis Functions')
para(doc, 'Testing Objective: Ensure that AR room analysis helper functions correctly process images and detect room properties.')
para(doc, 'Table 49: Unit Testing 7 AR Analysis', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Result'], [
    ['1', 'Detect dominant colors', 'room image with blue walls', 'Returns color: "Blue", hex: "#4A90E2"', 'Pass'],
    ['2', 'Estimate room dimensions', 'room image with visible walls', 'Returns width, height, depth estimates', 'Pass'],
    ['3', 'Detect room style', 'modern furniture in image', 'Returns style: "Modern", confidence: 85%', 'Pass'],
    ['4', 'Calculate suitability', 'furniture vs room dimensions', 'Returns score: 75–95%', 'Pass'],
    ['5', 'Invalid image format', 'file: not an image', 'Returns error: "Invalid image"', 'Pass']
])

h2(doc, 'Functional Testing')
para(doc, 'The functional testing will take place after the unit testing. In this functional testing, the functionality of each of the module is tested. This is to ensure that the system produced meets the specifications and requirements.')

h3(doc, 'Functional Testing 1: Upload Listing Images')
para(doc, 'Testing Objective: Ensure that uploaded listing images/files are accepted for supported formats and rejected otherwise.')
para(doc, 'Table 50: Functional Testing 1', style='Caption')
simple_table(doc, ['No.', 'Test case', 'Attribute and value', 'Expected result', 'Actual result', 'Result'], [
    ['1', 'Upload valid image', 'image.jpg', 'File accepted and uploaded to storage', 'Uploaded successfully', 'Pass'],
    ['2', 'Upload unsupported file type', 'file.txt', 'Won’t allow', 'Didn’t allowed', 'Pass'],
    ['3', 'Upload large file (>10MB)', 'Big image.webp', 'Display warning: "File size too large"', 'Warning shown', 'Pass']
])

h3(doc, 'Functional Testing 2: AI Price Prediction')
para(doc, 'Testing Objective: Verify that the price prediction engine returns valid prices and confidence for different listings.')
para(doc, 'Table 51: Functional Testing 2', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'Predict price for mobile device', 'Category: Mobile, Brand: Samsung, RAM: 8GB, Storage: 128GB', 'Display predicted price along with confidence range', 'Prediction displayed correctly', 'Pass'],
    ['2', 'Predict price for laptop', 'Category: Laptop, Brand: HP, Processor: Intel i5, RAM: 8GB', 'Display predicted price along with confidence range', 'Prediction displayed correctly', 'Pass'],
    ['3', 'Prediction with high confidence', 'Well-defined and complete specifications', 'Display narrow confidence range', 'High confidence shown correctly', 'Pass'],
    ['4', 'Auto-approval range display', 'Successful price prediction', 'Display green auto-approval price range', 'Auto-approval range displayed correctly', 'Pass'],
    ['5', 'Re-prediction on field change', 'User modifies product specifications', 'System automatically triggers a new price prediction', 'Auto-prediction working correctly', 'Pass']
])

h3(doc, 'Functional Testing 3: User & Admin Dashboard Analytics')
para(doc, 'Testing Objective: Ensure that dashboards correctly display metrics for users and admins.')
para(doc, 'Table 52: Functional Testing 3', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'User dashboard with listings', 'Valid user_id with existing listings', 'Display total listings, total views, and contact count', 'Metrics displayed correctly', 'Pass'],
    ['2', 'User dashboard with no listings', 'User with zero listings', 'Display message: "No data available"', 'Message displayed correctly', 'Pass'],
    ['3', 'Admin dashboard metrics', 'Complete system data available', 'Display total users, total listings, and active listings', 'Metrics displayed correctly', 'Pass'],
    ['4', 'Admin dashboard with empty database', 'No users or listings in system', 'Display zero values for all metrics', 'Zero metrics displayed correctly', 'Pass'],
    ['5', 'Admin dashboard top sellers', 'Users with multiple listings', 'Display top 5 sellers ranked by number of listings', 'Top sellers displayed correctly', 'Pass']
])

h3(doc, 'Functional Testing 4: Listing Creation & Validation')
para(doc, 'Testing Objective: Ensure that listing creation, duplicate prevention, and price validation work correctly.')
para(doc, 'Table 53: Functional Testing 4', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'Create listing with valid data', 'Title, description, price, category, images provided', 'Listing created successfully and user is redirected to dashboard', 'Listing created and visible on dashboard', 'Pass'],
    ['2', 'Create listing with duplicate title', 'Title already exists (case-insensitive)', 'System rejects submission with error', 'Duplicate title rejected correctly', 'Pass'],
    ['3', 'Create listing with duplicate images', 'Same image uploaded twice', 'System rejects submission with error', 'Duplicate image rejected', 'Pass'],
    ['4', 'Create listing with alphabetic price', 'Price entered as "five thousand"', 'System rejects submission with error', 'Alphabetic price rejected', 'Pass'],
    ['5', 'Create listing with valid images', '1–7 images, JPG/PNG format, each under 10MB', 'Images uploaded successfully', 'Images uploaded and displayed correctly', 'Pass']
])

h3(doc, 'Functional Testing 5: Admin Approval System')
para(doc, 'Testing Objective: Ensure that the administrator can correctly approve, reject, and manage pending listings')
para(doc, 'Table 54: Functional Testing 5', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'View pending listings', 'Admin logged in', 'Display all listings with status set to "Pending"', 'Pending listings displayed correctly', 'Pass'],
    ['2', 'Approve listing manually', 'Admin clicks Approve button', 'Listing status changes to "Active" and notification is sent to seller', 'Status changed to active successfully', 'Pass'],
    ['3', 'Reject listing manually', 'Admin clicks Reject and provides reason', 'Listing status changes to "Rejected" and rejection reason is saved', 'Listing rejected with reason saved', 'Pass'],
    ['4', 'Auto-approval for listings', 'Price within ±20% of predicted price', 'Listing is automatically approved without admin intervention', 'Auto-approved successfully', 'Pass'],
    ['5', 'Manual review for listings', 'Price outside ±20% of predicted price', 'Listing is sent to admin queue for manual review', 'Listing sent to pending review correctly', 'Pass']
])

h3(doc, 'Functional Testing 6: 3D AR (Augmented Reality) Features')
para(doc, 'Testing Objective: Ensure that 3D AR model upload, preview, customization, and mobile viewing work correctly for furniture listings.')
para(doc, 'Table 55: Functional Testing 6', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'Upload room photo for AR analysis', 'Valid JPG/PNG image, furniture listing', 'Room analyzed for dimensions, style, lighting; AI insights displayed', 'Room analysis working correctly', 'Pass'],
    ['2', 'AR canvas preview with furniture', 'Room photo uploaded', 'Canvas viewer opens; furniture rendered as 3D rectangles with proper perspective', 'Canvas preview working correctly', 'Pass'],
    ['3', 'AR customization (color/material)', 'Change furniture material', 'Furniture color/material updates in real-time on canvas', 'Customization working correctly', 'Pass'],
    ['4', 'Drag and drop furniture placement', 'Drag furniture objects on canvas', 'Furniture repositions with shadows and perspective maintained', 'Drag-and-drop working correctly', 'Pass'],
    ['5', 'Multiple view angles', 'Switch between front/3D/top views', 'Canvas updates to display different perspective angles', 'View switching working correctly', 'Pass']
])

h2(doc, 'Business Rules Testing')
para(doc, 'Testing Objective: Verify that business logic rules are correctly applied for listings, price predictions, and anomalies.')
para(doc, 'Table 56: Business Rules Testing Matrix', style='Caption')
simple_table(doc, ['Condition / Rule', 'R1', 'R2', 'R3', 'R4', 'R5'], [
    ['Listing Title Same?', 'Yes', 'No', 'No', 'No', 'No'],
    ['Price unusually high?', 'Yes', 'Yes', 'No', 'No', 'No'],
    ['Very old listing?', 'No', 'No', 'No', 'Yes', 'No'],
    ['Invalid Input / Data Error?', 'No', 'No', 'No', 'Yes', 'No'],
    ['Action', 'A1 (Flag duplicate)', 'A2 (Mark unusually high price)', 'A3 (Flag outdated listing)', 'A4 (Flag invalid data entry)', 'A5 (No anomaly)']
])

h2(doc, 'Integration Testing')
para(doc, 'Integration testing ensures that different modules work correctly when connected together.')

h3(doc, 'Integration Testing 1: User Registration - Email Verification - Login')
para(doc, 'Testing Objective: Verify that the user registration, authentication, and initial dashboard flow are correctly integrated.')
para(doc, 'Table 57: Integration Testing 1', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'User registers and account is created', 'Valid username, email, and password', 'User account is created in the database and a success message is displayed', 'Account created successfully', 'Pass'],
    ['2', 'User registers with duplicate username', 'Username already exists', 'System displays error message', 'Duplicate username rejected', 'Pass'],
    ['3', 'Newly registered user logs in', 'Valid newly registered credentials', 'User logs in successfully and is redirected to the user dashboard', 'Login working correctly', 'Pass'],
    ['4', 'Newly registered user views empty dashboard', 'New user with no listings', 'Dashboard displays "No listings yet" message', 'Empty dashboard state working correctly', 'Pass']
])

h3(doc, 'Integration Testing 2: Listing Upload - AI Price Prediction - Dashboard Update')
para(doc, 'Testing Objective: Verify the end-to-end flow from uploading a listing to predicting price and updating dashboards.')
para(doc, 'Table 58: Integration Testing 2', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'Full listing pipeline', 'User uploads a new listing', 'Data is saved successfully, price is predicted, and dashboards are updated', 'End-to-end process worked correctly', 'Pass'],
    ['2', 'Auto-approval triggered', 'Listing price within approved range', 'Price predicted → listing auto-approved → listing becomes active in feed', 'Auto-approval working correctly', 'Pass'],
    ['3', 'Manual review triggered', 'Listing price outside approved range', 'Price predicted → listing sent to admin → status set to pending', 'Manual review triggered successfully', 'Pass'],
    ['4', 'User dashboard updates', 'User creates a new listing', 'Listing count increases on user dashboard', 'User dashboard updated correctly', 'Pass'],
    ['5', 'Admin dashboard updates', 'New listing added to system', 'Total listings count increases in admin analytics', 'Admin dashboard updated correctly', 'Pass']
])

h3(doc, 'Integration Testing 3: Login - Dashboard - Listings')
para(doc, 'Table 59: Integration Testing 3', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'User logs in and dashboard loads', 'Valid user credentials', 'Dashboard loads successfully with listings feed', 'Dashboard loaded correctly', 'Pass'],
    ['2', 'User clicks Post Ad', 'Logged-in user', 'User is able to create and upload a listing', 'Listing upload works correctly', 'Pass'],
    ['3', 'Buyer-behavior browsing', 'User behaving as buyer only', 'User can browse and scroll listings without posting', 'Browsing works correctly', 'Pass'],
    ['4', 'User views listings', 'Any logged-in user', 'User can view all available listings', 'Listings displayed correctly', 'Pass']
])

h3(doc, 'Integration Testing 4: Category Selection - Form Fields - Price Prediction')
para(doc, 'Testing Objective: Verify that dynamic form fields update correctly based on category selection.')
para(doc, 'Table 60: Integration Testing 4', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'User selects Mobile category', 'Category set to Mobile', 'Brand, RAM, storage fields are displayed', 'Mobile-specific fields displayed correctly', 'Pass'],
    ['2', 'User enters mobile specifications', 'Brand: Apple, RAM: 8GB, Storage: 256GB', 'Predicted price is displayed along with confidence range', 'Price prediction working correctly', 'Pass'],
    ['3', 'User selects Laptop category', 'Category set to Laptop', 'Processor, GPU, RAM, storage fields displayed', 'Laptop-specific fields displayed correctly', 'Pass'],
    ['4', 'User changes category selection', 'Switch from Mobile to Laptop', 'Mobile fields are hidden and laptop fields are displayed', 'Fields updated correctly', 'Pass'],
    ['5', 'User completes specifications', 'All required fields filled', 'Green auto-approval price range is displayed', 'Auto-approval range displayed correctly', 'Pass']
])

h3(doc, 'Integration Testing 5: Favorites System - Dashboard - Listing Updates')
para(doc, 'Testing Objective: Verify that the favorites system is correctly integrated.')
para(doc, 'Table 61: Integration Testing 5', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'User adds a listing to favorites', 'User clicks favorite icon', 'Favorite is saved to database', 'Favorite saved successfully', 'Pass'],
    ['2', 'User adds favorite and dashboard updates', 'Listing added to favorites', 'Favorite count increases on user dashboard', 'Dashboard updated correctly', 'Pass'],
    ['3', 'User views favorites', 'User navigates to favorites page', 'All favorited listings are displayed', 'Favorites displayed correctly', 'Pass'],
    ['4', 'User removes a favorite', 'User clicks unfavorite icon', 'Favorite is removed, icon changes to unfilled state', 'Favorite removed successfully', 'Pass'],
    ['5', 'Listing deletion updates favorites', 'Seller deletes a listing', 'Listing is automatically removed from all users’ favorites', 'Favorites cleanup working correctly', 'Pass']
])

h3(doc, 'Integration Testing 6: 3D AR Model Upload → Storage → Preview')
para(doc, 'Testing Objective: Verify the end-to-end 3D AR workflow.')
para(doc, 'Table 62: Integration Testing 6', style='Caption')
simple_table(doc, ['No.', 'Test Case', 'Attribute and Value', 'Expected Result', 'Actual Result', 'Result'], [
    ['1', 'Upload 3D model and store on server', 'Upload GLB file for furniture', 'Model saved to ar_previews directory', 'Model stored successfully', 'Pass'],
    ['2', 'AR badge displayed', 'Furniture listing with 3D model', 'AR View badge visible on listing card', 'Badge displayed correctly', 'Pass'],
    ['3', 'User clicks AR preview', 'Click View in AR button', 'AR viewer opens; 3D model renders', 'AR viewer working correctly', 'Pass'],
    ['4', 'User customizes AR model', 'Change color/material and save', 'Customizations saved and reflected in all AR views', 'Customizations saved successfully', 'Pass']
])

# ══════════════════════════════════════════════════════════════════════════════
# CHAPTER 7 & 8 – CONCLUSION & REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1(doc, 'Conclusion and Future Work')
h2(doc, 'Conclusion')
para(doc, 'EzSell successfully addresses the core challenges of the informal Pakistani second-hand economy—specifically the "trust deficit" and "pricing anxiety"—by leveraging cutting-edge Artificial Intelligence. The integration of high-performance ML ensembles for price prediction, CLIP-based computer vision for fraud prevention, and Tripo AI for immersive AR visualization has transformed a traditionally fragmented process into a secure, structured, and data-driven digital marketplace. The project demonstrates that cognitive services can bridge the gap between buyer expectations and seller reliability, ultimately creating a more transparent and efficient ecosystem for used goods.')

h2(doc, 'Future Work')
bul(doc, 'In-App Escrow & Payments: Implementing a secure escrow mechanism and integrating local payment gateways like JazzCash and EasyPaisa to facilitate seamless, trust-based transactions.')
bul(doc, 'Monetization & Business Model Development: Designing a robust business model that introduces tiered services. This includes a "Freemium" approach where basic listings are free, while premium features—such as higher-visibility "Featured Ad" placement, AI-enhanced image editing, and advanced seller analytics—are available for a subscription or per-use fee.')
bul(doc, 'Dynamic Pricing Models: Introducing diverse pricing strategies tailored to different socio-economic segments, including subscription-based models for high-volume dealers and pay-as-you-go options for casual sellers to maximize platform accessibility.')
bul(doc, 'Mobile Ecosystem Expansion: Developing native iOS and Android applications using React Native or Flutter to provide a more responsive and accessible mobile experience.')
bul(doc, 'Categorical Scaling: Expanding the platform’s intelligence to complex high-value categories, including Vehicles and Real Estate, requiring specialized ML models for valuation.')

doc.add_page_break()
h1(doc, 'References')
refs = [
    'Radford, A., et al. (2021). "Learning Transferable Visual Models From Natural Language Supervision." (CLIP Architecture). International Conference on Machine Learning (ICML).',
    'Reimers, N., & Gurevych, I. (2019). "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks." arXiv preprint arXiv:1908.10084.',
    'Chen, T., & Guestrin, C. (2016). "XGBoost: A Scalable Tree Boosting System." Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining.',
    'Meta AI (2024). "Llama 3: Open Foundation and Fine-Tuned Chat Models." Meta AI Research & Groq LPU API Documentation.',
    'Ramírez, S. (2024). "FastAPI: A modern, high-performance web framework for building APIs with Python." Official Documentation.',
    'Meta Open Source (2024). "React: A JavaScript library for building user interfaces." Official React Documentation.',
    'Amazon Web Services (2024). "Amazon EC2 Instance Types and Cloud Infrastructure Best Practices." AWS Documentation.',
    'Tripo AI (2024). "Tripo AI V2 API: Automated Image-to-3D Model Generation and GLB Optimization." Technical Documentation.',
    'Supabase Team (2024). "Supabase: Open Source Realtime PostgreSQL and Storage Infrastructure." Documentation.',
    'Johnson, J., et al. (2019). "Billion-scale similarity search with GPUs." IEEE Transactions on Big Data.',
    'Jones, M., et al. (2015). "JSON Web Token (JWT) Standards." IETF RFC 7519.'
]
for r in refs: bul(doc, r)

doc.save(OUTPUT)
print(f'Master Document generation complete. Saved: {OUTPUT}')
